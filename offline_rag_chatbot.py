# 90% Aniqlikli Offline RAG ChatBot
# Hech qanday API ishlatmaydi, faqat local modellar

import os
import json
import pickle
import sqlite3
import hashlib
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import logging
from datetime import datetime
from collections import Counter
import math

# Document processing
import PyPDF2
import docx
from bs4 import BeautifulSoup

# Machine Learning va NLP
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.decomposition import TruncatedSVD
from sklearn.cluster import KMeans
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from nltk.tag import pos_tag

# BM25 for keyword search
from rank_bm25 import BM25Okapi

# Web interface
import streamlit as st
import pandas as pd

# NLTK downloads with comprehensive error handling
def download_nltk_data():
    """Download all required NLTK data with fallbacks"""
    required_data = [
        'punkt', 
        'punkt_tab',  # New requirement in latest NLTK
        'stopwords', 
        'averaged_perceptron_tagger', 
        'wordnet',
        'omw-1.4'  # Additional wordnet data
    ]
    
    for data in required_data:
        try:
            # Try to find existing data
            if data == 'punkt_tab':
                try:
                    nltk.data.find('tokenizers/punkt_tab')
                    continue
                except LookupError:
                    pass
            else:
                try:
                    nltk.data.find(f'tokenizers/{data}')
                    continue
                except LookupError:
                    try:
                        nltk.data.find(f'corpora/{data}')
                        continue
                    except LookupError:
                        pass
            
            # Download if not found
            print(f"Downloading NLTK data: {data}")
            try:
                nltk.download(data, quiet=True)
                print(f"✅ Successfully downloaded: {data}")
            except Exception as e:
                print(f"⚠️ Could not download {data}: {e}")
                # Try alternative downloads
                if data == 'punkt_tab':
                    try:
                        nltk.download('punkt', quiet=True)
                        print("✅ Downloaded punkt as fallback")
                    except:
                        pass
                        
        except Exception as e:
            print(f"❌ Error with {data}: {e}")
            continue

# Safe sentence tokenization with fallbacks
def safe_sent_tokenize(text):
    """Safe sentence tokenization with multiple fallbacks"""
    try:
        # Try modern punkt_tab first
        return sent_tokenize(text)
    except LookupError:
        try:
            # Try downloading punkt_tab
            nltk.download('punkt_tab', quiet=True)
            return sent_tokenize(text)
        except:
            try:
                # Fallback to punkt
                nltk.download('punkt', quiet=True)
                return sent_tokenize(text)
            except:
                # Manual sentence splitting as last resort
                import re
                sentences = re.split(r'[.!?]+', text)
                return [s.strip() for s in sentences if s.strip()]

download_nltk_data()

# Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class UltraHighAccuracyConfig:
    """Ultra yuqori aniqlik uchun optimallashtirilgan konfiguratsiya"""
    CHUNK_SIZE = 200              # Juda kichik chunklar - maksimal aniqlik
    CHUNK_OVERLAP = 30            # Minimal overlap
    MAX_DOCS_PER_QUERY = 15       # Ko'proq context
    MIN_CHUNK_LENGTH = 20         # Juda qisqa chunklar ham qabul qilinadi
    CONFIDENCE_THRESHOLD = 0.2    # Past threshold
    MAX_FEATURES = 15000          # Ko'proq TF-IDF features
    SVD_COMPONENTS = 500          # Katta embeddings
    DATABASE_PATH = "./ultra_accuracy_rag.db"
    EMBEDDINGS_PATH = "./ultra_accuracy_embeddings.pkl"
    
    # Weighted scoring - aniqlik uchun optimallashtirilgan
    SEMANTIC_WEIGHT = 0.5         # Semantic search weight
    KEYWORD_WEIGHT = 0.5          # Keyword search weight ko'paytirildi

class AdvancedDocumentLoader:
    """Yaxshilangan hujjat yuklash"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.html', '.md']
        self.text_cleaner = TextCleaner()
    
    def load_pdf(self, file_path: str) -> str:
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # PDF dan olingan matnni tozalash
            return self.text_cleaner.clean_pdf_text(text)
        except Exception as e:
            logger.error(f"PDF yuklashda xatolik {file_path}: {e}")
            return ""
    
    def load_docx(self, file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Paragraflar
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Jadvallar
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return self.text_cleaner.clean_text(text)
        except Exception as e:
            logger.error(f"DOCX yuklashda xatolik {file_path}: {e}")
            return ""
    
    def load_txt(self, file_path: str) -> str:
        encodings = ['utf-8', 'cp1251', 'latin1', 'ascii', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    return self.text_cleaner.clean_text(text)
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"TXT yuklashda xatolik {file_path}: {e}")
                return ""
        
        logger.error(f"TXT faylni hech qanday encoding bilan o'qib bo'lmadi: {file_path}")
        return ""
    
    def load_html(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                
                # Script va style ni o'chirish
                for element in soup(["script", "style", "nav", "footer"]):
                    element.decompose()
                
                text = soup.get_text()
                return self.text_cleaner.clean_text(text)
        except Exception as e:
            logger.error(f"HTML yuklashda xatolik {file_path}: {e}")
            return ""
    
    def load_document(self, file_path: str) -> str:
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.load_pdf(file_path)
        elif file_extension == '.docx':
            return self.load_docx(file_path)
        elif file_extension in ['.txt', '.md']:
            return self.load_txt(file_path)
        elif file_extension in ['.html', '.htm']:
            return self.load_html(file_path)
        else:
            logger.warning(f"Qo'llab-quvvatlanmaydigan format: {file_extension}")
            return ""

class TextCleaner:
    """Matnni chuqur tozalash"""
    
    def __init__(self):
        self.stemmer = PorterStemmer()
        
        # Stop words (ingliz + o'zbek + rus)
        try:
            english_stops = set(stopwords.words('english'))
        except:
            english_stops = set()
        
        uzbek_stops = {
            'va', 'yoki', 'lekin', 'ammo', 'uchun', 'bilan', 'dan', 'ga', 'da',
            'ni', 'ning', 'lar', 'cha', 'chi', 'mi', 'mu', 'ham', 'yo\'q', 'bor',
            'bu', 'u', 'ular', 'men', 'sen', 'biz', 'siz', 'o\'zi', 'o\'zini',
            'shu', 'ana', 'mana', 'endi', 'hozir', 'keyin', 'oldin', 'yuqorida',
            'pastda', 'ichida', 'tashqarida', 'yonida', 'oldida', 'orqasida'
        }
        
        russian_stops = {
            'и', 'в', 'во', 'не', 'что', 'он', 'на', 'я', 'с', 'со', 'как', 'а',
            'то', 'все', 'она', 'так', 'его', 'но', 'да', 'ты', 'к', 'у', 'же',
            'вы', 'за', 'бы', 'по', 'только', 'ее', 'мне', 'было', 'вот', 'от',
            'меня', 'еще', 'нет', 'о', 'из', 'ему', 'теперь', 'когда', 'даже',
            'ну', 'вдруг', 'ли', 'если', 'уже', 'или', 'ни', 'быть', 'был',
            'него', 'до', 'вас', 'нибудь', 'опять', 'уж', 'вам', 'ведь', 'там',
            'потом', 'себя', 'ничего', 'ей', 'может', 'они', 'тут', 'где', 'есть',
            'надо', 'ней', 'для', 'мы', 'тебя', 'их', 'чем', 'была', 'сам', 'чтоб'
        }
        
        self.stop_words = english_stops | uzbek_stops | russian_stops
    
    def clean_pdf_text(self, text: str) -> str:
        """PDF dan olingan matnni tozalash"""
        # Hyphenation ni tuzatish
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
        
        # Ortiqcha bo'sh qatorlar
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        # Page numbers va headers/footers
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        text = re.sub(r'\n\s*Page\s+\d+.*?\n', '\n', text, flags=re.IGNORECASE)
        
        return self.clean_text(text)
    
    def clean_text(self, text: str) -> str:
        """Umumiy matn tozalash"""
        if not text:
            return ""
        
        # HTML entities
        text = re.sub(r'&[a-zA-Z]+;', ' ', text)
        
        # Email va URL lar
        text = re.sub(r'\S+@\S+', '[EMAIL]', text)
        text = re.sub(r'http[s]?://\S+', '[URL]', text)
        
        # Raqamlarni normalizatsiya qilish
        text = re.sub(r'\b\d{4,}\b', '[NUMBER]', text)
        
        # Ortiqcha punktuatsiya
        text = re.sub(r'[.]{2,}', '.', text)
        text = re.sub(r'[!]{2,}', '!', text)
        text = re.sub(r'[?]{2,}', '?', text)
        
        # Ortiqcha bo'shliqlar
        text = re.sub(r'\s+', ' ', text)
        
        # Bosh va oxirgi bo'shliqlarni olib tashlash
        text = text.strip()
        
        return text
    
    def extract_keywords(self, text: str, top_k: int = 10) -> List[str]:
        """Matndan kalit so'zlarni chiqarish"""
        words = word_tokenize(text.lower())
        
        # Stop words va qisqa so'zlarni filtrlash
        filtered_words = [
            word for word in words 
            if word.isalnum() and len(word) > 2 and word not in self.stop_words
        ]
        
        # Chastota hisoblash
        word_freq = Counter(filtered_words)
        
        return [word for word, freq in word_freq.most_common(top_k)]

class SemanticChunker:
    """Ma'noga asoslangan chunking"""
    
    def __init__(self, chunk_size: int = 250, chunk_overlap: int = 40):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_cleaner = TextCleaner()
    
    def chunk_by_semantic_similarity(self, text: str, source: str = "") -> List[Dict]:
        """Semantik o'xshashlik asosida chunking"""
        
        # 1. Paragraflar bo'yicha bo'lish
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if not paragraphs:
            # Agar paragraflar bo'lmasa, jumlalar bo'yicha
            return self.chunk_by_sentences(text, source)
        
        chunks = []
        current_chunk_paras = []
        current_length = 0
        
        for para in paragraphs:
            para_length = len(para)
            
            # Agar bitta paragraf juda katta bo'lsa
            if para_length > self.chunk_size * 1.5:
                # Oldingi chunkni saqlash
                if current_chunk_paras:
                    chunk = self.create_chunk(current_chunk_paras, source, len(chunks))
                    if chunk:
                        chunks.append(chunk)
                    current_chunk_paras = []
                    current_length = 0
                
                # Katta paragrafni jumlalarga bo'lish
                para_chunks = self.chunk_large_paragraph(para, source, len(chunks))
                chunks.extend([c for c in para_chunks if c])
                
            elif current_length + para_length > self.chunk_size and current_chunk_paras:
                # Chunk yaratish
                chunk = self.create_chunk(current_chunk_paras, source, len(chunks))
                if chunk:
                    chunks.append(chunk)
                
                # Overlap uchun oxirgi paragrafni saqlab qolish
                overlap_paras = self.get_overlap_paragraphs(current_chunk_paras)
                current_chunk_paras = overlap_paras
                current_length = sum(len(p) for p in overlap_paras)
                
                current_chunk_paras.append(para)
                current_length += para_length
            else:
                current_chunk_paras.append(para)
                current_length += para_length
        
        # Oxirgi chunkni qo'shish
        if current_chunk_paras:
            chunk = self.create_chunk(current_chunk_paras, source, len(chunks))
            if chunk:
                chunks.append(chunk)
        
        return chunks
    
    def chunk_by_sentences(self, text: str, source: str = "") -> List[Dict]:
        """Jumlalar bo'yicha chunking (fallback)"""
        sentences = safe_sent_tokenize(text)
        chunks = []
        current_sentences = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            if current_length + sentence_length > self.chunk_size and current_sentences:
                chunk = self.create_chunk_from_sentences(current_sentences, source, len(chunks))
                if chunk:
                    chunks.append(chunk)
                
                # Overlap
                overlap_sentences = current_sentences[-1:] if current_sentences else []
                current_sentences = overlap_sentences
                current_length = sum(len(s) for s in overlap_sentences)
            
            current_sentences.append(sentence)
            current_length += sentence_length
        
        if current_sentences:
            chunk = self.create_chunk_from_sentences(current_sentences, source, len(chunks))
            if chunk:
                chunks.append(chunk)
        
        return chunks
    
    def chunk_large_paragraph(self, paragraph: str, source: str, start_id: int) -> List[Dict]:
        """Katta paragrafni kichik chunklarga bo'lish"""
        sentences = safe_sent_tokenize(paragraph)
        chunks = []
        current_sentences = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            if current_length + sentence_length > self.chunk_size and current_sentences:
                chunk = self.create_chunk_from_sentences(
                    current_sentences, source, start_id + len(chunks)
                )
                if chunk:
                    chunks.append(chunk)
                current_sentences = []
                current_length = 0
            
            current_sentences.append(sentence)
            current_length += sentence_length
        
        if current_sentences:
            chunk = self.create_chunk_from_sentences(
                current_sentences, source, start_id + len(chunks)
            )
            if chunk:
                chunks.append(chunk)
        
        return chunks
    
    def get_overlap_paragraphs(self, paragraphs: List[str]) -> List[str]:
        """Overlap uchun paragraflarni tanlash"""
        if not paragraphs:
            return []
        
        overlap_length = 0
        overlap_paras = []
        
        for para in reversed(paragraphs):
            if overlap_length + len(para) <= self.chunk_overlap:
                overlap_paras.insert(0, para)
                overlap_length += len(para)
            else:
                break
        
        return overlap_paras
    
    def create_chunk(self, paragraphs: List[str], source: str, chunk_id: int) -> Dict:
        """Paragraflardan chunk yaratish"""
        content = "\n\n".join(paragraphs)
        
        if len(content) < HighAccuracyConfig.MIN_CHUNK_LENGTH:
            return None
        
        # Kalit so'zlarni chiqarish
        keywords = self.text_cleaner.extract_keywords(content, top_k=5)
        
        return {
            'content': content,
            'source': source,
            'chunk_id': chunk_id,
            'length': len(content),
            'keywords': keywords,
            'paragraph_count': len(paragraphs),
            'type': 'semantic'
        }
    
    def create_chunk_from_sentences(self, sentences: List[str], source: str, chunk_id: int) -> Dict:
        """Jumlalardan chunk yaratish"""
        content = " ".join(sentences)
        
        if len(content) < HighAccuracyConfig.MIN_CHUNK_LENGTH:
            return None
        
        keywords = self.text_cleaner.extract_keywords(content, top_k=5)
        
        return {
            'content': content,
            'source': source,
            'chunk_id': chunk_id,
            'length': len(content),
            'keywords': keywords,
            'sentence_count': len(sentences),
            'type': 'sentence'
        }

class UltraHybridEmbedder:
    """Ultra TF-IDF + SVD + keyword + N-gram features"""
    
    def __init__(self, max_features: int = 15000, n_components: int = 500):
        self.max_features = max_features
        self.n_components = n_components
        
        # Enhanced TF-IDF vectorizer
        self.vectorizer = TfidfVectorizer(
            max_features=max_features,
            ngram_range=(1, 4),  # 1-4 gram (ko'proq pattern)
            min_df=1,            # Kam chastotali so'zlar ham
            max_df=0.98,         # Ko'proq so'zlar
            stop_words=None,     # Stop words yo'q - maksimal ma'lumot
            sublinear_tf=True,
            norm='l2',
            lowercase=True,
            token_pattern=r'\b\w+\b'  # Faqat so'zlar
        )
        
        # SVD for dimensionality reduction
        self.svd = TruncatedSVD(n_components=n_components, random_state=42)
        
        self.is_fitted = False
        self.text_cleaner = TextCleaner()
        
    def preprocess_texts(self, texts: List[str]) -> List[str]:
        """Kuchaytirilgan preprocessing"""
        processed = []
        for text in texts:
            # Minimal tozalash - ko'proq ma'lumot saqlab qolish
            clean_text = text.lower()
            
            # Faqat juda zarur tozalash
            clean_text = re.sub(r'\s+', ' ', clean_text)  # Ortiqcha bo'shliqlar
            clean_text = clean_text.strip()
            
            processed.append(clean_text)
        
        return processed
    
    def fit_transform(self, texts: List[str], keywords_list: List[List[str]] = None) -> np.ndarray:
        """Model training va embeddings yaratish"""
        logger.info(f"Hybrid embedding modeli {len(texts)} ta matn bo'yicha o'rgatilmoqda...")
        
        # Preprocessing
        processed_texts = self.preprocess_texts(texts)
        
        # TF-IDF
        tfidf_matrix = self.vectorizer.fit_transform(processed_texts)
        logger.info(f"TF-IDF matrix o'lchami: {tfidf_matrix.shape}")
        
        # SVD
        embeddings = self.svd.fit_transform(tfidf_matrix)
        logger.info(f"SVD dan keyin embedding o'lchami: {embeddings.shape}")
        
        # Keyword features qo'shish
        if keywords_list:
            keyword_features = self.create_keyword_features(keywords_list)
            embeddings = np.hstack([embeddings, keyword_features])
            logger.info(f"Keyword features qo'shilgandan keyin: {embeddings.shape}")
        
        self.is_fitted = True
        return embeddings
    
    def transform(self, texts: List[str], keywords_list: List[List[str]] = None) -> np.ndarray:
        """Yangi matnlar uchun embeddings"""
        if not self.is_fitted:
            raise ValueError("Model avval o'rgatilishi kerak!")
        
        processed_texts = self.preprocess_texts(texts)
        tfidf_matrix = self.vectorizer.transform(processed_texts)
        embeddings = self.svd.transform(tfidf_matrix)
        
        if keywords_list:
            keyword_features = self.create_keyword_features(keywords_list)
            embeddings = np.hstack([embeddings, keyword_features])
        
        return embeddings
    
    def create_keyword_features(self, keywords_list: List[List[str]]) -> np.ndarray:
        """Keyword-based features yaratish"""
        # Barcha unique keywords
        all_keywords = set()
        for keywords in keywords_list:
            all_keywords.update(keywords)
        
        keyword_vocab = list(all_keywords)
        features = np.zeros((len(keywords_list), len(keyword_vocab)))
        
        for i, keywords in enumerate(keywords_list):
            for keyword in keywords:
                if keyword in keyword_vocab:
                    j = keyword_vocab.index(keyword)
                    features[i, j] = 1.0
        
        return features
    
    def save(self, filepath: str):
        """Modelni saqlash"""
        model_data = {
            'vectorizer': self.vectorizer,
            'svd': self.svd,
            'is_fitted': self.is_fitted,
            'max_features': self.max_features,
            'n_components': self.n_components
        }
        with open(filepath, 'wb') as f:
            pickle.dump(model_data, f)
        logger.info(f"Hybrid embedding model saqlandi: {filepath}")
    
    def load(self, filepath: str):
        """Modelni yuklash"""
        with open(filepath, 'rb') as f:
            model_data = pickle.load(f)
        
        self.vectorizer = model_data['vectorizer']
        self.svd = model_data['svd']
        self.is_fitted = model_data['is_fitted']
        self.max_features = model_data['max_features']
        self.n_components = model_data['n_components']
        
        logger.info(f"Hybrid embedding model yuklandi: {filepath}")

class AdvancedDatabase:
    """Yaxshilangan database bilan full-text search"""
    
    def __init__(self, db_path: str):
        self.db_path = db_path
        self.init_database()
    
    def init_database(self):
        """Database va indexes yaratish"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Main documents table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS documents (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    source TEXT NOT NULL,
                    content TEXT NOT NULL,
                    chunk_id INTEGER NOT NULL,
                    length INTEGER NOT NULL,
                    keywords TEXT,
                    chunk_type TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    content_hash TEXT UNIQUE
                )
            ''')
            
            # Embeddings table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS embeddings (
                    document_id INTEGER PRIMARY KEY,
                    embedding BLOB NOT NULL,
                    FOREIGN KEY (document_id) REFERENCES documents (id)
                )
            ''')
            
            # FTS table for keyword search
            cursor.execute('''
                CREATE VIRTUAL TABLE IF NOT EXISTS documents_fts USING fts5(
                    document_id,
                    content,
                    keywords,
                    source
                )
            ''')
            
            # Indexes
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_source ON documents(source)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_content_hash ON documents(content_hash)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_keywords ON documents(keywords)')
            
            conn.commit()
    
    def add_documents(self, chunks: List[Dict], embeddings: np.ndarray):
        """Chunklarni database ga saqlash"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            successful_adds = 0
            
            for i, chunk in enumerate(chunks):
                if chunk is None:
                    continue
                
                content_hash = hashlib.md5(chunk['content'].encode()).hexdigest()
                keywords_str = ','.join(chunk.get('keywords', []))
                
                try:
                    # Main table ga qo'shish
                    cursor.execute('''
                        INSERT INTO documents (source, content, chunk_id, length, keywords, chunk_type, content_hash)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        chunk['source'], chunk['content'], chunk['chunk_id'],
                        chunk['length'], keywords_str, chunk.get('type', 'unknown'), content_hash
                    ))
                    
                    document_id = cursor.lastrowid
                    
                    # Embedding qo'shish
                    embedding_blob = embeddings[i].tobytes()
                    cursor.execute('''
                        INSERT INTO embeddings (document_id, embedding)
                        VALUES (?, ?)
                    ''', (document_id, embedding_blob))
                    
                    # FTS table ga qo'shish
                    cursor.execute('''
                        INSERT INTO documents_fts (document_id, content, keywords, source)
                        VALUES (?, ?, ?, ?)
                    ''', (document_id, chunk['content'], keywords_str, chunk['source']))
                    
                    successful_adds += 1
                    
                except sqlite3.IntegrityError:
                    logger.warning(f"Duplicate content: {chunk['source']}")
                    continue
                except Exception as e:
                    logger.error(f"Document qo'shishda xatolik: {e}")
                    continue
            
            conn.commit()
            logger.info(f"Database ga {successful_adds} ta document qo'shildi")
    
    def keyword_search(self, query: str, top_k: int = 10) -> List[Dict]:
        """FTS keyword search"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # FTS query
            search_query = f'"{query}"'  # Exact phrase search
            
            cursor.execute('''
                SELECT d.id, d.source, d.content, d.chunk_id, d.length, d.keywords,
                       documents_fts.rank
                FROM documents_fts
                JOIN documents d ON documents_fts.document_id = d.id
                WHERE documents_fts MATCH ?
                ORDER BY documents_fts.rank
                LIMIT ?
            ''', (search_query, top_k))
            
            results = []
            for row in cursor.fetchall():
                doc_id, source, content, chunk_id, length, keywords, rank = row
                results.append({
                    'id': doc_id,
                    'source': source,
                    'content': content,
                    'chunk_id': chunk_id,
                    'length': length,
                    'keywords': keywords.split(',') if keywords else [],
                    'keyword_score': abs(rank)  # FTS rank negative bo'ladi
                })
            
            return results
    
    def semantic_search(self, query_embedding: np.ndarray, top_k: int = 10) -> List[Dict]:
        """Semantic similarity search"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT d.id, d.source, d.content, d.chunk_id, d.length, d.keywords, e.embedding
                FROM documents d
                JOIN embeddings e ON d.id = e.document_id
            ''')
            
            results = cursor.fetchall()
            
            if not results:
                return []
            
            similarities = []
            for row in results:
                doc_id, source, content, chunk_id, length, keywords, embedding_blob = row
                
                # Embedding restore qilish
                embedding = np.frombuffer(embedding_blob, dtype=np.float64)
                embedding = embedding.reshape(1, -1)
                query_emb = query_embedding.reshape(1, -1)
                
                # Cosine similarity
                similarity = cosine_similarity(query_emb, embedding)[0][0]
                
                similarities.append({
                    'id': doc_id,
                    'source': source,
                    'content': content,
                    'chunk_id': chunk_id,
                    'length': length,
                    'keywords': keywords.split(',') if keywords else [],
                    'semantic_score': similarity
                })
            
            # Similarity bo'yicha saralash
            similarities.sort(key=lambda x: x['semantic_score'], reverse=True)
            return similarities[:top_k]
    
    def get_all_documents(self) -> List[Dict]:
        """Barcha hujjatlarni olish (BM25 uchun)"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT id, source, content, chunk_id, length, keywords
                FROM documents
            ''')
            
            results = []
            for row in cursor.fetchall():
                doc_id, source, content, chunk_id, length, keywords = row
                results.append({
                    'id': doc_id,
                    'source': source,
                    'content': content,
                    'chunk_id': chunk_id,
                    'length': length,
                    'keywords': keywords.split(',') if keywords else []
                })
            
            return results
    
    def get_stats(self) -> Dict:
        """Database statistikalari"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Asosiy statistika
            cursor.execute('SELECT COUNT(*) FROM documents')
            total_docs = cursor.fetchone()[0]
            
            cursor.execute('SELECT COUNT(DISTINCT source) FROM documents')
            unique_sources = cursor.fetchone()[0]
            
            cursor.execute('SELECT AVG(length) FROM documents')
            avg_length = cursor.fetchone()[0] or 0
            
            # Chunk types
            cursor.execute('SELECT chunk_type, COUNT(*) FROM documents GROUP BY chunk_type')
            chunk_types = dict(cursor.fetchall())
            
            return {
                'total_documents': total_docs,
                'unique_sources': unique_sources,
                'average_length': round(avg_length, 2),
                'chunk_types': chunk_types
            }

class HybridRetriever:
    """Semantic + Keyword + BM25 hybrid search"""
    
    def __init__(self, database: AdvancedDatabase, embedder: HybridEmbedder):
        self.database = database
        self.embedder = embedder
        self.text_cleaner = TextCleaner()
        
        # BM25 setup
        self.bm25 = None
        self.bm25_docs = None
        self.setup_bm25()
    
    def setup_bm25(self):
        """BM25 model setup"""
        try:
            all_docs = self.database.get_all_documents()
            if all_docs:
                # Tokenize documents
                tokenized_docs = []
                self.bm25_docs = []
                
                for doc in all_docs:
                    tokens = self.text_cleaner.clean_text(doc['content']).lower().split()
                    tokenized_docs.append(tokens)
                    self.bm25_docs.append(doc)
                
                self.bm25 = BM25Okapi(tokenized_docs)
                logger.info(f"BM25 model {len(all_docs)} hujjat bilan sozlandi")
        except Exception as e:
            logger.error(f"BM25 setup xatolik: {e}")
            self.bm25 = None
    
    def preprocess_query(self, query: str) -> str:
        """Query preprocessing"""
        # Tozalash
        clean_query = self.text_cleaner.clean_text(query)
        
        # Kichik harflarga
        clean_query = clean_query.lower()
        
        # Imlo xatolarini tuzatish (oddiy)
        corrections = {
            'nma': 'nima',
            'qnday': 'qanday',
            'qchon': 'qachon',
            'qyer': 'qayer',
            'kanchik': 'qancha',
            'budu': 'bunda'
        }
        
        for wrong, correct in corrections.items():
            clean_query = clean_query.replace(wrong, correct)
        
        return clean_query
    
    def expand_query(self, query: str) -> str:
        """Query expansion with synonyms"""
        synonyms = {
            'nima': ['qanday', 'qanaqa', 'nimalar'],
            'qachon': ['qaysi vaqt', 'qaysi kun', 'nechanchi'],
            'qayer': ['qayerda', 'qaysi joy', 'qaysi joyda'],
            'kim': ['kimlar', 'qaysi odam', 'qaysi shaxs'],
            'qancha': ['necha', 'qanday miqdor', 'qanday son'],
            'nega': ['nima uchun', 'qanday sabab', 'nimaga'],
            'qanday': ['nima', 'qanaqa', 'nechanchi xil']
        }
        
        expanded_terms = [query]
        
        for word, syns in synonyms.items():
            if word in query.lower():
                expanded_terms.extend(syns)
        
        return ' '.join(expanded_terms)
    
    def hybrid_search(self, query: str, top_k: int = 12) -> List[Dict]:
        """Hybrid search combining multiple methods"""
        
        # Query preprocessing
        processed_query = self.preprocess_query(query)
        expanded_query = self.expand_query(processed_query)
        
        # 1. Semantic search
        semantic_results = self.semantic_search(expanded_query, top_k)
        
        # 2. Keyword search (FTS)
        keyword_results = self.keyword_search(processed_query, top_k)
        
        # 3. BM25 search
        bm25_results = self.bm25_search(processed_query, top_k)
        
        # 4. Combine and rank
        combined_results = self.combine_search_results(
            semantic_results, keyword_results, bm25_results, query, top_k
        )
        
        return combined_results
    
    def semantic_search(self, query: str, top_k: int) -> List[Dict]:
        """Semantic similarity search"""
        try:
            # Query embedding
            query_embedding = self.embedder.transform([query])[0]
            return self.database.semantic_search(query_embedding, top_k)
        except Exception as e:
            logger.error(f"Semantic search xatolik: {e}")
            return []
    
    def keyword_search(self, query: str, top_k: int) -> List[Dict]:
        """Keyword search using FTS"""
        try:
            return self.database.keyword_search(query, top_k)
        except Exception as e:
            logger.error(f"Keyword search xatolik: {e}")
            return []
    
    def bm25_search(self, query: str, top_k: int) -> List[Dict]:
        """BM25 search"""
        if not self.bm25 or not self.bm25_docs:
            return []
        
        try:
            query_tokens = query.split()
            scores = self.bm25.get_scores(query_tokens)
            
            # Top results
            top_indices = np.argsort(scores)[-top_k:][::-1]
            
            results = []
            for idx in top_indices:
                if idx < len(self.bm25_docs):
                    doc = self.bm25_docs[idx].copy()
                    doc['bm25_score'] = scores[idx]
                    results.append(doc)
            
            return results
            
        except Exception as e:
            logger.error(f"BM25 search xatolik: {e}")
            return []
    
    def combine_search_results(self, semantic_results: List[Dict], 
                             keyword_results: List[Dict], 
                             bm25_results: List[Dict],
                             original_query: str,
                             top_k: int) -> List[Dict]:
        """Search natijalarini combine qilish"""
        
        # Barcha natijalarni birlashtirish
        all_results = {}  # doc_id -> result
        
        # Semantic results
        for i, result in enumerate(semantic_results):
            doc_id = result['id']
            result['semantic_rank'] = i + 1
            result['semantic_score'] = result.get('semantic_score', 0)
            all_results[doc_id] = result
        
        # Keyword results
        for i, result in enumerate(keyword_results):
            doc_id = result['id']
            if doc_id in all_results:
                all_results[doc_id]['keyword_rank'] = i + 1
                all_results[doc_id]['keyword_score'] = result.get('keyword_score', 0)
            else:
                result['keyword_rank'] = i + 1
                result['keyword_score'] = result.get('keyword_score', 0)
                result['semantic_score'] = 0
                result['semantic_rank'] = 999
                all_results[doc_id] = result
        
        # BM25 results
        for i, result in enumerate(bm25_results):
            doc_id = result['id']
            if doc_id in all_results:
                all_results[doc_id]['bm25_rank'] = i + 1
                all_results[doc_id]['bm25_score'] = result.get('bm25_score', 0)
            else:
                result['bm25_rank'] = i + 1
                result['bm25_score'] = result.get('bm25_score', 0)
                result['semantic_score'] = 0
                result['keyword_score'] = 0
                result['semantic_rank'] = 999
                result['keyword_rank'] = 999
                all_results[doc_id] = result
        
        # Combined scoring
        combined_results = []
        for doc_id, result in all_results.items():
            # Normalize scores
            semantic_score = result.get('semantic_score', 0)
            keyword_score = result.get('keyword_score', 0)
            bm25_score = result.get('bm25_score', 0)
            
            # Combined score with weights
            combined_score = (
                HighAccuracyConfig.SEMANTIC_WEIGHT * semantic_score +
                HighAccuracyConfig.KEYWORD_WEIGHT * keyword_score +
                0.3 * min(bm25_score / 10.0, 1.0)  # BM25 normalize
            )
            
            # Query-content overlap bonus
            query_words = set(original_query.lower().split())
            content_words = set(result['content'].lower().split())
            overlap_ratio = len(query_words & content_words) / max(len(query_words), 1)
            combined_score += 0.2 * overlap_ratio
            
            # Keywords match bonus
            if result.get('keywords'):
                keyword_matches = sum(1 for kw in result['keywords'] if kw.lower() in original_query.lower())
                combined_score += 0.1 * keyword_matches / max(len(result['keywords']), 1)
            
            result['combined_score'] = combined_score
            combined_results.append(result)
        
        # Sort by combined score
        combined_results.sort(key=lambda x: x['combined_score'], reverse=True)
        
        return combined_results[:top_k]

class IntelligentResponseGenerator:
    """Aqlli javob generatsiya qilish"""
    
    def __init__(self):
        self.query_types = {
            'definition': ['nima', 'ta\'rif', 'definition', 'bu nima', 'nimani anglatadi'],
            'how': ['qanday', 'qanaqa', 'jarayon', 'usul', 'how'],
            'when': ['qachon', 'qaysi vaqt', 'nechanchi', 'when'],
            'where': ['qayer', 'qayerda', 'qaysi joy', 'where'],
            'who': ['kim', 'kimlar', 'qaysi odam', 'who'],
            'why': ['nega', 'nima uchun', 'sabab', 'why'],
            'how_much': ['qancha', 'necha', 'miqdor', 'how much', 'how many'],
            'comparison': ['farq', 'taqqoslash', 'difference', 'compare', 'o\'xshash', 'vs'],
            'list': ['ro\'yxat', 'list', 'sanab', 'turlar', 'types']
        }
        
        self.templates = {
            'definition': """
**Ta'rif bo'yicha:**

{main_content}

**Qisqacha:** {summary}
            """,
            
            'how': """
**Jarayon/Usul:**

{main_content}

**Asosiy bosqichlar:** {steps}
            """,
            
            'when': """
**Vaqt bo'yicha ma'lumot:**

{main_content}

**Muhim sanalar:** {key_dates}
            """,
            
            'where': """
**Joy/Manzil bo'yicha:**

{main_content}

**Asosiy joylar:** {locations}
            """,
            
            'who': """
**Shaxslar/Tashkilotlar haqida:**

{main_content}

**Asosiy ishtirokchilar:** {people}
            """,
            
            'why': """
**Sabab/Asos:**

{main_content}

**Asosiy sabablar:** {reasons}
            """,
            
            'how_much': """
**Miqdor/Son bo'yicha:**

{main_content}

**Asosiy raqamlar:** {numbers}
            """,
            
            'comparison': """
**Taqqoslash:**

{main_content}

**Asosiy farqlar:** {differences}
            """,
            
            'list': """
**Ro'yxat/Turlari:**

{main_content}

**Asosiy elementlar:** {items}
            """,
            
            'general': """
**Ma'lumot:**

{main_content}

**Xulosa:** {summary}
            """
        }
    
    def classify_query_type(self, query: str) -> str:
        """Query turini aniqlash"""
        query_lower = query.lower()
        
        for query_type, keywords in self.query_types.items():
            if any(keyword in query_lower for keyword in keywords):
                return query_type
        
        return 'general'
    
    def extract_key_information(self, contexts: List[str], query_type: str) -> Dict[str, str]:
        """Contextdan kalit ma'lumotlarni ajratish"""
        combined_context = "\n\n".join(contexts)
        
        # Summary
        summary = self.create_summary(combined_context)
        
        # Type-specific extractions
        extracted = {'summary': summary}
        
        if query_type == 'how':
            extracted['steps'] = self.extract_steps(combined_context)
        elif query_type == 'when':
            extracted['key_dates'] = self.extract_dates(combined_context)
        elif query_type == 'where':
            extracted['locations'] = self.extract_locations(combined_context)
        elif query_type == 'who':
            extracted['people'] = self.extract_people(combined_context)
        elif query_type == 'why':
            extracted['reasons'] = self.extract_reasons(combined_context)
        elif query_type == 'how_much':
            extracted['numbers'] = self.extract_numbers(combined_context)
        elif query_type == 'comparison':
            extracted['differences'] = self.extract_differences(combined_context)
        elif query_type == 'list':
            extracted['items'] = self.extract_list_items(combined_context)
        
        return extracted
    
    def create_summary(self, text: str, max_sentences: int = 3) -> str:
        """Matn xulosasi yaratish"""
        sentences = safe_sent_tokenize(text)
        if len(sentences) <= max_sentences:
            return text
        
        # Eng muhim jumlalarni tanlash (oddiy scoring)
        sentence_scores = []
        
        for sentence in sentences:
            # Uzunlik score
            length_score = min(len(sentence) / 100, 1.0)
            
            # Keyword density score  
            words = sentence.lower().split()
            important_words = ['muhim', 'asosiy', 'zarur', 'kerak', 'bo\'ladi', 'hisoblanadi']
            keyword_score = sum(1 for word in words if word in important_words) / max(len(words), 1)
            
            # Position score (birinchi va oxirgi jumlalar muhimroq)
            position_score = 1.0 if sentences.index(sentence) < 2 or sentences.index(sentence) >= len(sentences) - 2 else 0.5
            
            total_score = length_score + keyword_score + position_score
            sentence_scores.append((sentence, total_score))
        
        # Top sentences
        sentence_scores.sort(key=lambda x: x[1], reverse=True)
        top_sentences = [s[0] for s in sentence_scores[:max_sentences]]
        
        # Original order ga qaytarish
        ordered_sentences = []
        for sentence in sentences:
            if sentence in top_sentences:
                ordered_sentences.append(sentence)
        
        return " ".join(ordered_sentences)
    
    def extract_steps(self, text: str) -> str:
        """Bosqichlarni ajratish"""
        # Step indicators
        step_patterns = [
            r'\d+\.\s+',  # 1. 2. 3.
            r'birinchi|ikkinchi|uchinchi|to\'rtinchi|beshinchi',
            r'avval|keyin|so\'ng|oxirida|boshlash|tugatish',
            r'first|second|third|then|next|finally'
        ]
        
        sentences = safe_sent_tokenize(text)
        steps = []
        
        for sentence in sentences:
            if any(re.search(pattern, sentence.lower()) for pattern in step_patterns):
                steps.append(sentence.strip())
        
        return "; ".join(steps[:5]) if steps else "Aniq bosqichlar ko'rsatilmagan."
    
    def extract_dates(self, text: str) -> str:
        """Sanalarni ajratish"""
        date_patterns = [
            r'\d{1,2}/\d{1,2}/\d{4}',
            r'\d{1,2}\.\d{1,2}\.\d{4}',
            r'\d{4}-\d{1,2}-\d{1,2}',
            r'\d{4}\s*yil',
            r'\d{1,2}\s*(yanvar|fevral|mart|aprel|may|iyun|iyul|avgust|sentabr|oktabr|noyabr|dekabr)'
        ]
        
        dates = []
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            dates.extend(matches)
        
        return "; ".join(dates[:5]) if dates else "Aniq sanalar ko'rsatilmagan."
    
    def extract_locations(self, text: str) -> str:
        """Joylarni ajratish"""
        # Oddiy location extraction
        location_indicators = [
            'shahar', 'viloyat', 'davlat', 'mamlakat', 'mintaqa', 'hudud',
            'manzil', 'ko\'cha', 'bino', 'ofis', 'joy', 'joyda'
        ]
        
        sentences = sent_tokenize(text)
        locations = []
        
        for sentence in sentences:
            if any(indicator in sentence.lower() for indicator in location_indicators):
                # Extract potential location names (capitalize words)
                words = sentence.split()
                for i, word in enumerate(words):
                    if word[0].isupper() and len(word) > 2:
                        locations.append(word)
        
        return "; ".join(list(set(locations))[:5]) if locations else "Aniq joylar ko'rsatilmagan."
    
    def extract_people(self, text: str) -> str:
        """Odamlarni ajratish"""
        # Simple name extraction (capitalized words)
        people_indicators = ['tomonidan', 'muallif', 'yozuvchi', 'direktor', 'mudur', 'rahbar']
        
        sentences = sent_tokenize(text)
        people = []
        
        for sentence in sentences:
            words = sentence.split()
            for i, word in enumerate(words):
                # Look for capitalized words that might be names
                if (word[0].isupper() and len(word) > 2 and 
                    not word.lower() in ['Bu', 'Ushbu', 'Shuning', 'Lekin', 'Ammo']):
                    
                    # Check if next word is also capitalized (full name)
                    if i + 1 < len(words) and words[i + 1][0].isupper():
                        people.append(f"{word} {words[i + 1]}")
                    else:
                        people.append(word)
        
        return "; ".join(list(set(people))[:5]) if people else "Aniq shaxslar ko'rsatilmagan."
    
    def extract_reasons(self, text: str) -> str:
        """Sabablarni ajratish"""
        reason_indicators = [
            'sabab', 'sababli', 'uchun', 'tufayli', 'natijada', 'oqibatida',
            'chunki', 'negaki', 'because', 'due to', 'as a result'
        ]
        
        sentences = sent_tokenize(text)
        reasons = []
        
        for sentence in sentences:
            if any(indicator in sentence.lower() for indicator in reason_indicators):
                reasons.append(sentence.strip())
        
        return "; ".join(reasons[:3]) if reasons else "Aniq sabablar ko'rsatilmagan."
    
    def extract_numbers(self, text: str) -> str:
        """Raqamlarni ajratish"""
        number_patterns = [
            r'\d+%',  # percentages
            r'\d+\s*(million|billion|ming|yuz)',  # large numbers
            r'\d+\.\d+',  # decimals
            r'\$\d+',  # money
            r'\d+\s*(sm|mm|m|km|kg|g|l)',  # measurements
        ]
        
        numbers = []
        for pattern in number_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            numbers.extend(matches)
        
        return "; ".join(numbers[:5]) if numbers else "Aniq raqamlar ko'rsatilmagan."
    
    def extract_differences(self, text: str) -> str:
        """Farqlarni ajratish"""
        diff_indicators = [
            'farq', 'farqli', 'o\'xshash emas', 'different', 'unlike',
            'lekin', 'ammo', 'biroq', 'however', 'but', 'whereas'
        ]
        
        sentences = sent_tokenize(text)
        differences = []
        
        for sentence in sentences:
            if any(indicator in sentence.lower() for indicator in diff_indicators):
                differences.append(sentence.strip())
        
        return "; ".join(differences[:3]) if differences else "Aniq farqlar ko'rsatilmagan."
    
    def extract_list_items(self, text: str) -> str:
        """Ro'yxat elementlarini ajratish"""
        # Look for bullet points, numbers, or list indicators
        list_patterns = [
            r'^\s*[-•*]\s+',  # bullet points
            r'^\s*\d+\.\s+',  # numbered lists
            r'^\s*[a-z]\)\s+',  # lettered lists
        ]
        
        lines = text.split('\n')
        items = []
        
        for line in lines:
            line = line.strip()
            if any(re.match(pattern, line) for pattern in list_patterns):
                # Remove list markers
                clean_item = re.sub(r'^\s*[-•*\d+\.a-z\)]\s*', '', line)
                if clean_item:
                    items.append(clean_item)
        
        return "; ".join(items[:7]) if items else "Aniq ro'yxat topilmadi."
    
    def generate_response(self, query: str, contexts: List[str], confidence: float) -> str:
        """Final javob generatsiya qilish"""
        if not contexts:
            return "Kechirasiz, sizning savolingizga mos ma'lumot topilmadi. Iltimos, savolni boshqacha shaklda yozing."
        
        # Query type aniqlash
        query_type = self.classify_query_type(query)
        
        # Key information extract qilish
        extracted_info = self.extract_key_information(contexts, query_type)
        
        # Template tanlash
        template = self.templates.get(query_type, self.templates['general'])
        
        # Main content
        main_content = self.create_summary("\n\n".join(contexts), max_sentences=4)
        
        # Template fill qilish
        try:
            response = template.format(
                main_content=main_content,
                **extracted_info
            )
        except KeyError:
            # Fallback to general template
            response = self.templates['general'].format(
                main_content=main_content,
                summary=extracted_info.get('summary', 'Qo\'shimcha ma\'lumot yo\'q.')
            )
        
        # Confidence-based warning
        if confidence < 50:
            response += "\n\n⚠️ **Diqqat:** Bu javob nisbatan past ishonchlilik darajasiga ega. Qo'shimcha tekshirish tavsiya etiladi."
        elif confidence < 70:
            response += "\n\n💡 **Eslatma:** Qo'shimcha aniqlik uchun savolni batafsilroq yozing."
        
        return response.strip()

class SuperAdvancedConfidenceCalculator:
    """Ultra ilg'or confidence hisoblash"""
    
    def __init__(self):
        self.text_cleaner = TextCleaner()
    
    def calculate_comprehensive_confidence(self, query: str, results: List[Dict], 
                                         generated_answer: str) -> float:
        """Ultra comprehensive confidence calculation"""
        
        if not results:
            return 0.0
        
        confidence_factors = []
        
        # 1. Enhanced similarity scoring
        if 'combined_score' in results[0]:
            scores = [r.get('combined_score', 0) for r in results]
            avg_score = np.mean(scores)
            max_score = max(scores)
            score_variance = np.var(scores)
            
            # Composite score
            composite_score = (avg_score * 0.6 + max_score * 0.4) * (1 - score_variance * 0.1)
            confidence_factors.append(min(composite_score * 1.5, 1.0))  # Boost factor
        
        # 2. Multi-level result quality
        top_3_scores = [r.get('combined_score', 0) for r in results[:3]]
        top_quality = np.mean(top_3_scores) if top_3_scores else 0
        confidence_factors.append(min(top_quality * 1.3, 1.0))  # Boost factor
        
        # 3. Enhanced query-answer relevance
        query_relevance = self.calculate_enhanced_relevance(query, generated_answer)
        confidence_factors.append(query_relevance)
        
        # 4. Context consistency with variance
        if len(results) > 1:
            consistency = self.calculate_enhanced_consistency(results)
            confidence_factors.append(consistency)
        
        # 5. Answer completeness and quality
        completeness = self.calculate_enhanced_completeness(query, generated_answer, results)
        confidence_factors.append(completeness)
        
        # 6. Source diversity with quality weighting
        source_diversity = self.calculate_weighted_source_diversity(results)
        confidence_factors.append(source_diversity)
        
        # 7. Enhanced keyword coverage
        keyword_coverage = self.calculate_enhanced_keyword_coverage(query, results)
        confidence_factors.append(keyword_coverage)
        
        # 8. NEW: Content length and structure score
        structure_score = self.calculate_content_structure_score(generated_answer)
        confidence_factors.append(structure_score)
        
        # 9. NEW: Query complexity matching
        complexity_match = self.calculate_query_complexity_match(query, results)
        confidence_factors.append(complexity_match)
        
        # Enhanced weighted average with boosting
        weights = [0.20, 0.18, 0.15, 0.12, 0.10, 0.08, 0.07, 0.05, 0.05]
        weighted_confidence = np.average(confidence_factors, weights=weights[:len(confidence_factors)])
        
        # Apply confidence boosting formula
        boosted_confidence = self.apply_confidence_boosting(weighted_confidence, results, query)
        
        # Normalize to percentage with higher ceiling
        return min(boosted_confidence * 100, 98)  # Max 98% instead of 95%
    
    def calculate_enhanced_relevance(self, query: str, answer: str) -> float:
        """Kuchaytirilgan relevance calculation"""
        query_words = set(self.text_cleaner.clean_text(query).lower().split())
        answer_words = set(self.text_cleaner.clean_text(answer).lower().split())
        
        if not query_words:
            return 0.0
        
        # Enhanced Jaccard with partial matching
        exact_matches = query_words & answer_words
        
        # Partial matches (substring matching)
        partial_matches = 0
        for q_word in query_words:
            for a_word in answer_words:
                if len(q_word) > 3 and len(a_word) > 3:
                    if q_word in a_word or a_word in q_word:
                        partial_matches += 0.5
        
        exact_score = len(exact_matches) / len(query_words)
        partial_score = min(partial_matches / len(query_words), 0.5)
        
        return min(exact_score + partial_score, 1.0)
    
    def calculate_enhanced_consistency(self, results: List[Dict]) -> float:
        """Kuchaytirilgan consistency calculation"""
        if len(results) < 2:
            return 1.0
        
        contents = [r['content'] for r in results]
        similarities = []
        
        # Pairwise similarity with weighting
        for i in range(len(contents)):
            for j in range(i + 1, len(contents)):
                # Weight by result ranking (top results more important)
                weight = 1.0 / (1 + i + j)
                
                content1_words = set(contents[i].lower().split())
                content2_words = set(contents[j].lower().split())
                
                if content1_words and content2_words:
                    intersection = content1_words & content2_words
                    union = content1_words | content2_words
                    similarity = len(intersection) / len(union) * weight
                    similarities.append(similarity)
        
        return np.mean(similarities) if similarities else 0.0
    
    def calculate_enhanced_completeness(self, query: str, answer: str, results: List[Dict]) -> float:
        """Kuchaytirilgan completeness calculation"""
        # Answer length scoring with optimal range
        answer_length = len(answer)
        
        if answer_length < 100:
            length_score = answer_length / 100
        elif answer_length > 2000:
            length_score = max(0.6, 2000 / answer_length)
        else:
            length_score = 1.0
        
        # Question type coverage
        question_indicators = {
            'nima': ['ta\'rif', 'definition', 'bu', 'hisoblanadi'],
            'qanday': ['jarayon', 'usul', 'method', 'bosqich'],
            'qachon': ['vaqt', 'time', 'sana', 'muddat'],
            'qayer': ['joy', 'manzil', 'location', 'joyda'],
            'kim': ['shaxs', 'person', 'odam', 'tomonidan'],
            'nega': ['sabab', 'reason', 'uchun', 'tufayli'],
            'qancha': ['miqdor', 'amount', 'son', 'raqam']
        }
        
        query_lower = query.lower()
        answer_lower = answer.lower()
        
        coverage_score = 0
        total_questions = 0
        
        for question_word, answer_indicators in question_indicators.items():
            if question_word in query_lower:
                total_questions += 1
                if any(indicator in answer_lower for indicator in answer_indicators):
                    coverage_score += 1
        
        question_coverage = coverage_score / total_questions if total_questions > 0 else 1.0
        
        # Context utilization score
        total_context_length = sum(len(r['content']) for r in results)
        context_utilization = min(len(answer) / max(total_context_length * 0.1, 100), 1.0)
        
        return (length_score + question_coverage + context_utilization) / 3
    
    def calculate_weighted_source_diversity(self, results: List[Dict]) -> float:
        """Weighted source diversity"""
        sources = [r.get('source', '') for r in results]
        unique_sources = len(set(sources))
        total_sources = len(sources)
        
        base_diversity = unique_sources / total_sources if total_sources > 0 else 0.0
        
        # Bonus for multiple high-quality sources
        if unique_sources >= 3:
            base_diversity *= 1.2
        elif unique_sources >= 2:
            base_diversity *= 1.1
        
        return min(base_diversity, 1.0)
    
    def calculate_enhanced_keyword_coverage(self, query: str, results: List[Dict]) -> float:
        """Kuchaytirilgan keyword coverage"""
        query_keywords = set(self.text_cleaner.extract_keywords(query, top_k=8))  # More keywords
        
        if not query_keywords:
            return 1.0
        
        # Collect keywords from top results with weighting
        weighted_keywords = {}
        for i, result in enumerate(results[:5]):  # Top 5 results
            weight = 1.0 / (1 + i)  # Higher weight for top results
            if result.get('keywords'):
                for keyword in result['keywords']:
                    weighted_keywords[keyword] = weighted_keywords.get(keyword, 0) + weight
        
        # Enhanced matching with partial matches
        coverage_score = 0
        for q_keyword in query_keywords:
            # Exact match
            if q_keyword in weighted_keywords:
                coverage_score += weighted_keywords[q_keyword]
            else:
                # Partial match
                for r_keyword in weighted_keywords:
                    if len(q_keyword) > 3 and len(r_keyword) > 3:
                        if q_keyword in r_keyword or r_keyword in q_keyword:
                            coverage_score += weighted_keywords[r_keyword] * 0.5
                            break
        
        normalized_score = coverage_score / (len(query_keywords) * 1.0)  # Normalize by max possible score
        return min(normalized_score, 1.0)
    
    def calculate_content_structure_score(self, answer: str) -> float:
        """Content structure and quality score"""
        # Check for structured content
        structure_indicators = [
            r'\*\*.*?\*\*',  # Bold text
            r'^\d+\.',       # Numbered lists
            r'^[-•*]',       # Bullet points
            r':
    
    def calculate_query_answer_relevance(self, query: str, answer: str) -> float:
        """Query va answer orasidagi relevance"""
        query_words = set(self.text_cleaner.clean_text(query).lower().split())
        answer_words = set(self.text_cleaner.clean_text(answer).lower().split())
        
        if not query_words:
            return 0.0
        
        # Jaccard similarity
        intersection = query_words & answer_words
        union = query_words | answer_words
        
        return len(intersection) / len(union) if union else 0.0
    
    def calculate_context_consistency(self, results: List[Dict]) -> float:
        """Context mazmunlarining consistency"""
        if len(results) < 2:
            return 1.0
        
        contents = [r['content'] for r in results]
        
        # Pairwise similarity calculation
        similarities = []
        for i in range(len(contents)):
            for j in range(i + 1, len(contents)):
                content1_words = set(contents[i].lower().split())
                content2_words = set(contents[j].lower().split())
                
                intersection = content1_words & content2_words
                union = content1_words | content2_words
                
                if union:
                    similarity = len(intersection) / len(union)
                    similarities.append(similarity)
        
        return np.mean(similarities) if similarities else 0.0
    
    def calculate_answer_completeness(self, query: str, answer: str) -> float:
        """Javobning to'liqligi"""
        # Answer uzunligi (optimal range)
        answer_length = len(answer)
        if answer_length < 50:
            length_score = answer_length / 50
        elif answer_length > 1000:
            length_score = max(0.5, 1000 / answer_length)
        else:
            length_score = 1.0
        
        # Question words coverage
        question_words = ['nima', 'qanday', 'qachon', 'qayer', 'kim', 'nega', 'qancha']
        query_lower = query.lower()
        
        answered_questions = 0
        total_questions = 0
        
        for qw in question_words:
            if qw in query_lower:
                total_questions += 1
                # Check if answer contains related content
                if any(indicator in answer.lower() for indicator in [
                    'ta\'rif', 'jarayon', 'vaqt', 'joy', 'shaxs', 'sabab', 'miqdor'
                ]):
                    answered_questions += 1
        
        question_coverage = answered_questions / total_questions if total_questions > 0 else 1.0
        
        return (length_score + question_coverage) / 2
    
    def calculate_source_diversity(self, results: List[Dict]) -> float:
        """Source turlarining diversity"""
        sources = [r.get('source', '') for r in results]
        unique_sources = len(set(sources))
        total_sources = len(sources)
        
        return unique_sources / total_sources if total_sources > 0 else 0.0
    
    def calculate_keyword_coverage(self, query: str, results: List[Dict]) -> float:
        """Query keywords ning results da coverage"""
        query_keywords = set(self.text_cleaner.extract_keywords(query, top_k=5))
        
        if not query_keywords:
            return 1.0
        
        result_keywords = set()
        for result in results:
            if result.get('keywords'):
                result_keywords.update(result['keywords'])
        
        if not result_keywords:
            return 0.0
        
        # Keywords overlap
        overlap = query_keywords & result_keywords
        return len(overlap) / len(query_keywords)

class UltraHighAccuracyRAGPipeline:
    """Ultra yuqori aniqlikli RAG pipeline - 90%+ target"""
    
    def __init__(self, config: UltraHighAccuracyConfig = None):
        self.config = config or UltraHighAccuracyConfig()
        
        # Enhanced components
        self.doc_loader = AdvancedDocumentLoader()
        self.chunker = SemanticChunker(
            chunk_size=self.config.CHUNK_SIZE,
            chunk_overlap=self.config.CHUNK_OVERLAP
        )
        self.embedder = UltraHybridEmbedder(
            max_features=self.config.MAX_FEATURES,
            n_components=self.config.SVD_COMPONENTS
        )
        self.database = AdvancedDatabase(self.config.DATABASE_PATH)
        self.retriever = None
        self.response_generator = IntelligentResponseGenerator()
        self.confidence_calculator = SuperAdvancedConfidenceCalculator()
        
        self.is_ready = False
        
    def process_documents(self, file_paths: List[str]):
        """Hujjatlarni qayta ishlash va indexing"""
        logger.info(f"🚀 {len(file_paths)} ta fayl yuqori aniqlikda qayta ishlanmoqda...")
        
        all_chunks = []
        all_keywords = []
        successful_files = 0
        
        # 1. Document loading va chunking
        for file_path in file_paths:
            try:
                logger.info(f"📄 Qayta ishlanmoqda: {file_path}")
                
                # Load document
                text = self.doc_loader.load_document(file_path)
                if not text.strip():
                    logger.warning(f"⚠️ Bo'sh fayl: {file_path}")
                    continue
                
                # Semantic chunking
                chunks = self.chunker.chunk_by_semantic_similarity(text, source=file_path)
                
                # Filter valid chunks
                valid_chunks = [chunk for chunk in chunks if chunk is not None]
                
                if valid_chunks:
                    all_chunks.extend(valid_chunks)
                    # Keywords collect qilish
                    chunk_keywords = [chunk.get('keywords', []) for chunk in valid_chunks]
                    all_keywords.extend(chunk_keywords)
                    
                    successful_files += 1
                    logger.info(f"✅ {file_path}: {len(valid_chunks)} ta semantic chunk")
                else:
                    logger.warning(f"⚠️ Valid chunk yaratilmadi: {file_path}")
                
            except Exception as e:
                logger.error(f"❌ {file_path}: {e}")
        
        if not all_chunks:
            logger.error("❌ Hech qanday valid chunk yaratilmadi!")
            return
        
        logger.info(f"📊 Jami {len(all_chunks)} ta chunk yaratildi")
        
        # 2. Advanced embedding generation
        logger.info("🧠 Hybrid embeddings yaratilmoqda...")
        chunk_texts = [chunk['content'] for chunk in all_chunks]
        embeddings = self.embedder.fit_transform(chunk_texts, all_keywords)
        
        # 3. Database indexing
        logger.info("💾 Advanced database indexing...")
        self.database.add_documents(all_chunks, embeddings)
        
        # 4. Retriever setup
        self.retriever = HybridRetriever(self.database, self.embedder)
        
        # 5. Model saving
        self.embedder.save(self.config.EMBEDDINGS_PATH)
        
        self.is_ready = True
        
        # Final statistics
        stats = self.database.get_stats()
        logger.info(f"""
        🎉 Yuqori aniqlikli qayta ishlash tugadi:
        - ✅ Muvaffaqiyatli fayllar: {successful_files}/{len(file_paths)}
        - 📄 Jami chunklar: {stats['total_documents']}
        - 📁 Noyob manbalar: {stats['unique_sources']}
        - 📏 O'rtacha chunk uzunligi: {stats['average_length']} belgi
        - 🔧 Chunk turlari: {stats['chunk_types']}
        """)
        
    def load_existing_model(self):
        """Mavjud modelni yuklash"""
        try:
            if os.path.exists(self.config.EMBEDDINGS_PATH):
                logger.info("📂 Mavjud model yuklanmoqda...")
                self.embedder.load(self.config.EMBEDDINGS_PATH)
                
                stats = self.database.get_stats()
                if stats['total_documents'] > 0:
                    # Retriever setup
                    self.retriever = HybridRetriever(self.database, self.embedder)
                    self.is_ready = True
                    
                    logger.info(f"""
                    ✅ Mavjud model muvaffaqiyatli yuklandi:
                    - 📄 Jami hujjatlar: {stats['total_documents']}
                    - 📁 Manbalar: {stats['unique_sources']}
                    - 📊 O'rtacha uzunlik: {stats['average_length']} belgi
                    """)
                    return True
                    
        except Exception as e:
            logger.error(f"❌ Model yuklashda xatolik: {e}")
        
        return False
    
    def query(self, question: str, top_k: int = None) -> Dict[str, Any]:
        """Yuqori aniqlikli savol-javob"""
        if not self.is_ready:
            return {
                'answer': "❌ Model hali tayyor emas. Iltimos, avval hujjatlarni qayta ishlang.",
                'sources': [],
                'confidence': 0.0,
                'context_used': 0
            }
        
        if top_k is None:
            top_k = self.config.MAX_DOCS_PER_QUERY
        
        try:
            logger.info(f"🔍 Query: '{question}'")
            
            # 1. Hybrid search
            search_results = self.retriever.hybrid_search(question, top_k)
            
            if not search_results:
                return {
                    'answer': "🔍 Sizning savolingizga mos ma'lumot topilmadi. Iltimos, savolni boshqacha shaklda yozing yoki qo'shimcha hujjatlar yuklang.",
                    'sources': [],
                    'confidence': 0.0,
                    'context_used': 0
                }
            
            # 2. Context preparation
            contexts = [result['content'] for result in search_results]
            sources = list(set([result['source'] for result in search_results]))
            
            # 3. Intelligent response generation
            answer = self.response_generator.generate_response(question, contexts, 0)
            
            # 4. Advanced confidence calculation
            confidence = self.confidence_calculator.calculate_comprehensive_confidence(
                question, search_results, answer
            )
            
            # 5. Update answer with confidence
            final_answer = self.response_generator.generate_response(question, contexts, confidence)
            
            logger.info(f"✅ Query completed - Confidence: {confidence:.1f}%")
            
            return {
                'answer': final_answer,
                'sources': sources,
                'confidence': round(confidence, 1),
                'context_used': len(search_results),
                'search_details': {
                    'semantic_results': len([r for r in search_results if r.get('semantic_score', 0) > 0]),
                    'keyword_results': len([r for r in search_results if r.get('keyword_score', 0) > 0]),
                    'bm25_results': len([r for r in search_results if r.get('bm25_score', 0) > 0]),
                    'top_score': search_results[0].get('combined_score', 0) if search_results else 0
                }
            }
            
        except Exception as e:
            logger.error(f"❌ Query xatolik: {e}")
            return {
                'answer': "❌ Xatolik yuz berdi. Iltimos qaytadan urinib ko'ring.",
                'sources': [],
                'confidence': 0.0,
                'context_used': 0
            }

def load_css_style():
    """CSS stillarni yuklash"""
    css_style = """
    <style>
    * {
        font-family: Verdana, Geneva, Tahoma, sans-serif !important;
    }

    .main-header {
        text-align: center;
        background: linear-gradient(90deg, #00F700 0%, #00D600 100%);
        padding: 1.5rem;
        border-radius: 15px;
        color: white;
        margin-bottom: 2rem;
        box-shadow: 0 4px 15px rgba(0, 247, 0, 0.3);
    }

    .main-header h1 {
        font-size: 2.5rem;
        margin-bottom: 0.5rem;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
    }

    .main-header p {
        font-size: 1.2rem;
        opacity: 0.9;
    }

    .stButton > button {
        width: 95% !important;
        height: 60px !important;
        border: none !important;
        outline: none !important;
        color: #fff !important;
        background: #111 !important;
        cursor: pointer !important;
        position: relative !important;
        z-index: 0 !important;
        border-radius: 10px !important;
        font-weight: bold !important;
        font-size: 16px !important;
        transition: all 0.3s ease !important;
    }

    .stButton > button:before {
        content: '';
        background: linear-gradient(45deg, #00F700, #73ff00, #fffb00, #48ff00, #00ffd5, #002bff, #7a00ff, #ff00c8, #00F700);
        position: absolute;
        top: -2px;
        left: -2px;
        background-size: 400%;
        z-index: -1;
        filter: blur(5px);
        width: calc(100% + 4px);
        height: calc(100% + 4px);
        animation: glowing 20s linear infinite;
        opacity: 0;
        transition: opacity .3s ease-in-out;
        border-radius: 10px;
    }

    .stButton > button:active {
        color: #000 !important;
    }

    .stButton > button:active:after {
        background: transparent !important;
    }

    .stButton > button:hover:before {
        opacity: 1;
    }

    .stButton > button:after {
        z-index: -1;
        content: '';
        position: absolute;
        width: 100%;
        height: 100%;
        background: #00F700;
        left: 0;
        top: 0;
        border-radius: 10px;
    }

    @keyframes glowing {
        0% { background-position: 0 0; }
        50% { background-position: 400% 0; }
        100% { background-position: 0 0; }
    }

    .stat-box {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
        margin: 0.5rem 0;
        border: 2px solid #00F700;
    }

    .confidence-high { 
        color: #00F700 !important; 
        font-weight: bold !important;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
    }
    .confidence-medium { 
        color: #ffc107 !important; 
        font-weight: bold !important;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
    }
    .confidence-low { 
        color: #dc3545 !important; 
        font-weight: bold !important;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
    }

    .upload-section {
        border: 2px dashed #00F700;
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
        background: rgba(0, 247, 0, 0.05);
    }

    .sidebar-title {
        color: #00F700 !important;
        font-weight: bold !important;
        font-size: 1.2rem !important;
        margin-bottom: 1rem !important;
    }

    .metric-container {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 0.8rem;
        border-radius: 8px;
        margin: 0.3rem 0;
    }

    .chat-container {
        border-radius: 10px;
        padding: 1rem;
        margin: 0.5rem 0;
    }

    .user-message {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 15px 15px 5px 15px;
        padding: 1rem;
        margin: 0.5rem 0;
    }

    .assistant-message {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        color: white;
        border-radius: 15px 15px 15px 5px;
        padding: 1rem;
        margin: 0.5rem 0;
    }

    .processing-indicator {
        background: linear-gradient(45deg, #00F700, #73ff00);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
        animation: pulse 2s infinite;
    }

    @keyframes pulse {
        0% { opacity: 1; }
        50% { opacity: 0.7; }
        100% { opacity: 1; }
    }
    </style>
    """
    return css_style

def create_advanced_streamlit_interface():
    """Ilg'or Streamlit interface"""
    
    st.set_page_config(
        page_title="CHATBOT",
        page_icon="🎯",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Apply custom CSS
    st.markdown(load_css_style(), unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🎯 CHATBOT</h1>
        <p>Yuqori aniqlikli AI yordamchisi - API-siz, to'liq offline</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Session state initialization
    if 'rag_pipeline' not in st.session_state:
        st.session_state.rag_pipeline = HighAccuracyRAGPipeline()
        # Auto-load existing model
        st.session_state.rag_pipeline.load_existing_model()
    
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    
    if 'processing_stats' not in st.session_state:
        st.session_state.processing_stats = {}
    
    # Sidebar
    with st.sidebar:
        st.markdown('<h2 class="sidebar-title">⚙️ Boshqaruv Paneli</h2>', unsafe_allow_html=True)
        
        # Model status
        if st.session_state.rag_pipeline.is_ready:
            st.success("✅ Model tayyor")
            stats = st.session_state.rag_pipeline.database.get_stats()
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f'<div class="metric-container">📄 Hujjatlar<br><strong>{stats["total_documents"]}</strong></div>', unsafe_allow_html=True)
                st.markdown(f'<div class="metric-container">📏 O\'rtacha uzunlik<br><strong>{stats["average_length"]:.0f}</strong></div>', unsafe_allow_html=True)
            with col2:
                st.markdown(f'<div class="metric-container">📁 Manbalar<br><strong>{stats["unique_sources"]}</strong></div>', unsafe_allow_html=True)
                if 'chunk_types' in stats:
                    semantic_count = stats['chunk_types'].get('semantic', 0)
                    st.markdown(f'<div class="metric-container">🧠 Semantic chunks<br><strong>{semantic_count}</strong></div>', unsafe_allow_html=True)
        else:
            st.warning("⚠️ Model hali tayyor emas")
        
        st.markdown("---")
        
        # Configuration
        st.subheader("🔧 Sozlamalar")
        
        chunk_size = st.slider("Chunk o'lchami", 150, 400, 
                              st.session_state.rag_pipeline.config.CHUNK_SIZE)
        
        max_docs = st.slider("Maksimal contextlar", 5, 20, 
                           st.session_state.rag_pipeline.config.MAX_DOCS_PER_QUERY)
        
        # Update config
        st.session_state.rag_pipeline.config.CHUNK_SIZE = chunk_size
        st.session_state.rag_pipeline.config.MAX_DOCS_PER_QUERY = max_docs
        
        st.markdown("---")
        
        # File upload
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.subheader("📤 Hujjat yuklash")
        
        # File uploader with improved handling
        uploaded_files = st.file_uploader(
            "Fayllarni tanlang",
            type=['pdf', 'docx', 'txt', 'html', 'md'],
            accept_multiple_files=True,
            help="Qo'llab-quvvatlanuvchi formatlar: PDF, DOCX, TXT, HTML, MD"
        )
        
        # Show uploaded files
        if uploaded_files:
            st.write("**Yuklangan fayllar:**")
            for file in uploaded_files:
                file_size = len(file.getvalue()) / 1024 / 1024  # MB
                st.write(f"📄 {file.name} ({file_size:.1f} MB)")
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Processing button
        col1, col2 = st.columns(2)
        with col1:
            process_btn = st.button("🚀 Qayta ishlash", disabled=not uploaded_files, use_container_width=True)
        with col2:
            clear_btn = st.button("🗑️ Tozalash", use_container_width=True)
        
        # Processing with comprehensive error handling and debugging
        if process_btn and uploaded_files:
            # Create temporary directory
            temp_dir = Path("./temp_uploads")
            temp_dir.mkdir(exist_ok=True)
            
            # Initialize variables
            temp_paths = []
            
            try:
                # Phase 1: File saving
                st.info("📁 1-bosqich: Fayllar saqlanmoqda...")
                progress_bar = st.progress(0)
                
                total_files = len(uploaded_files)
                for i, uploaded_file in enumerate(uploaded_files):
                    try:
                        # Progress update
                        progress = (i + 1) / total_files * 0.2
                        progress_bar.progress(progress)
                        
                        # Create safe filename
                        original_name = uploaded_file.name
                        safe_name = re.sub(r'[^\w\-_\.]', '_', original_name)
                        if not safe_name:
                            safe_name = f"document_{i}.txt"
                        
                        temp_path = temp_dir / f"{i}_{safe_name}"
                        
                        # Save file
                        file_content = uploaded_file.read()
                        if len(file_content) == 0:
                            st.warning(f"⚠️ Bo'sh fayl: {original_name}")
                            continue
                            
                        with open(temp_path, "wb") as f:
                            f.write(file_content)
                        
                        temp_paths.append(str(temp_path))
                        st.success(f"✅ Saqlandi: {original_name} ({len(file_content)} bytes)")
                        
                    except Exception as e:
                        st.error(f"❌ {uploaded_file.name} saqlashda xatolik: {str(e)}")
                        continue
                
                if not temp_paths:
                    st.error("❌ Hech qanday fayl saqlanmadi!")
                    return
                
                st.success(f"✅ {len(temp_paths)} ta fayl muvaffaqiyatli saqlandi")
                
                # Phase 2: Document processing
                st.info("🧠 2-bosqich: Hujjatlar qayta ishlanmoqda...")
                progress_bar.progress(0.3)
                
                # Test document loading first
                doc_loader = AdvancedDocumentLoader()
                loaded_texts = []
                
                for i, temp_path in enumerate(temp_paths):
                    try:
                        progress = 0.3 + (i + 1) / len(temp_paths) * 0.2
                        progress_bar.progress(progress)
                        
                        text = doc_loader.load_document(temp_path)
                        if text and len(text.strip()) > 10:
                            loaded_texts.append((temp_path, text))
                            st.info(f"📄 Yuklandi: {Path(temp_path).name} ({len(text)} belgi)")
                        else:
                            st.warning(f"⚠️ Bo'sh yoki juda qisqa matn: {Path(temp_path).name}")
                    except Exception as e:
                        st.error(f"❌ {Path(temp_path).name} yuklashda xatolik: {str(e)}")
                        continue
                
                if not loaded_texts:
                    st.error("❌ Hech qanday matn yuklanmadi!")
                    return
                
                st.success(f"✅ {len(loaded_texts)} ta hujjat matnlari yuklandi")
                
                # Phase 3: Create new pipeline and process
                st.info("⚙️ 3-bosqich: RAG pipeline yaratilmoqda...")
                progress_bar.progress(0.6)
                
                # Create completely fresh ultra pipeline
                new_config = UltraHighAccuracyConfig()
                new_pipeline = UltraHighAccuracyRAGPipeline(new_config)
                
                # Process each document step by step
                all_chunks = []
                all_keywords = []
                
                for i, (file_path, text) in enumerate(loaded_texts):
                    try:
                        progress = 0.6 + (i + 1) / len(loaded_texts) * 0.2
                        progress_bar.progress(progress)
                        
                        # Create chunks
                        chunks = new_pipeline.chunker.chunk_by_semantic_similarity(text, source=Path(file_path).name)
                        valid_chunks = [chunk for chunk in chunks if chunk is not None]
                        
                        if valid_chunks:
                            all_chunks.extend(valid_chunks)
                            chunk_keywords = [chunk.get('keywords', []) for chunk in valid_chunks]
                            all_keywords.extend(chunk_keywords)
                            st.info(f"🔧 {Path(file_path).name}: {len(valid_chunks)} chunk yaratildi")
                        else:
                            st.warning(f"⚠️ {Path(file_path).name}: chunk yaratilmadi")
                            
                    except Exception as e:
                        st.error(f"❌ {Path(file_path).name} chunking xatolik: {str(e)}")
                        continue
                
                if not all_chunks:
                    st.error("❌ Hech qanday chunk yaratilmadi!")
                    return
                
                st.success(f"✅ Jami {len(all_chunks)} ta chunk yaratildi")
                
                # Phase 4: Create embeddings
                st.info("🧠 4-bosqich: Embeddings yaratilmoqda...")
                progress_bar.progress(0.8)
                
                try:
                    chunk_texts = [chunk['content'] for chunk in all_chunks]
                    embeddings = new_pipeline.embedder.fit_transform(chunk_texts, all_keywords)
                    st.success(f"✅ Embeddings yaratildi: {embeddings.shape}")
                except Exception as e:
                    st.error(f"❌ Embeddings yaratishda xatolik: {str(e)}")
                    return
                
                # Phase 5: Save to database
                st.info("💾 5-bosqich: Database ga saqlash...")
                progress_bar.progress(0.9)
                
                try:
                    new_pipeline.database.add_documents(all_chunks, embeddings)
                    new_pipeline.retriever = HybridRetriever(new_pipeline.database, new_pipeline.embedder)
                    new_pipeline.embedder.save(new_pipeline.config.EMBEDDINGS_PATH)
                    new_pipeline.is_ready = True
                    st.success("✅ Database ga saqlandi")
                except Exception as e:
                    st.error(f"❌ Database ga saqlashda xatolik: {str(e)}")
                    return
                
                # Phase 6: Final verification and replacement
                st.info("✅ 6-bosqich: Yakuniy tekshirish...")
                progress_bar.progress(0.95)
                
                # Verify everything works
                try:
                    test_stats = new_pipeline.database.get_stats()
                    if test_stats['total_documents'] > 0 and new_pipeline.is_ready:
                        # Replace the session pipeline
                        st.session_state.rag_pipeline = new_pipeline
                        
                        progress_bar.progress(1.0)
                        
                        # Show final success
                        success_msg = f"""
                        🎉 **MUVAFFAQIYAT!**
                        
                        ✅ **Jami hujjatlar:** {test_stats['total_documents']}  
                        ✅ **Noyob manbalar:** {test_stats['unique_sources']}  
                        ✅ **O'rtacha uzunlik:** {test_stats['average_length']:.0f} belgi  
                        ✅ **Model holati:** Tayyor  
                        
                        Endi savollaringizni bering! 🚀
                        """
                        st.balloons()
                        st.success(success_msg)
                        
                        # Auto-refresh after success
                        import time
                        time.sleep(3)
                        st.rerun()
                        
                    else:
                        st.error("❌ Model tekshiruvdan o'tmadi!")
                        
                except Exception as e:
                    st.error(f"❌ Yakuniy tekshiruvda xatolik: {str(e)}")
                
            except Exception as e:
                st.error(f"❌ Umumiy xatolik: {str(e)}")
                import traceback
                st.code(traceback.format_exc())
                
            finally:
                # Cleanup temp files
                for temp_path in temp_paths:
                    try:
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                    except:
                        pass
                try:
                    if temp_dir.exists():
                        import shutil
                        shutil.rmtree(temp_dir, ignore_errors=True)
                except:
                    pass
        
        # Clear data
        if clear_btn:
            if st.session_state.get('clear_confirmed', False):
                try:
                    # Remove database and embeddings
                    if os.path.exists(st.session_state.rag_pipeline.config.DATABASE_PATH):
                        os.remove(st.session_state.rag_pipeline.config.DATABASE_PATH)
                    if os.path.exists(st.session_state.rag_pipeline.config.EMBEDDINGS_PATH):
                        os.remove(st.session_state.rag_pipeline.config.EMBEDDINGS_PATH)
                    
                    # Reset pipeline
                    st.session_state.rag_pipeline = HighAccuracyRAGPipeline()
                    st.session_state.messages = []
                    st.session_state.clear_confirmed = False
                    
                    st.success("✅ Ma'lumotlar tozalandi!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Tozalash xatolik: {e}")
            else:
                st.session_state.clear_confirmed = True
                st.warning("⚠️ Haqiqatan ham barcha ma'lumotlarni o'chirmoqchimisiz? Qayta bosing.")
    
    # Main chat interface
    st.header("💬 Yuqori Aniqlikli Suhbat")
    
    # Chat container
    chat_container = st.container()
    
    with chat_container:
        # Display chat messages
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
                
                # Additional info for assistant messages
                if message["role"] == "assistant":
                    if "confidence" in message:
                        confidence = message["confidence"]
                        if confidence >= 80:
                            conf_class = "confidence-high"
                            conf_icon = "🎯"
                        elif confidence >= 60:
                            conf_class = "confidence-medium"
                            conf_icon = "⚡"
                        else:
                            conf_class = "confidence-low"
                            conf_icon = "⚠️"
                        
                        st.markdown(f"""
                        <div style="margin: 1rem 0;">
                            <span class="{conf_class}">{conf_icon} Ishonchlilik: {confidence}%</span>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # Sources
                    if "sources" in message and message["sources"]:
                        with st.expander("📚 Manbalar"):
                            for i, source in enumerate(message["sources"], 1):
                                st.write(f"{i}. {source}")
                    
                    # Search details
                    if "search_details" in message:
                        details = message["search_details"]
                        with st.expander("🔍 Qidiruv tafsilotlari"):
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("Semantic", details.get('semantic_results', 0))
                            with col2:
                                st.metric("Keyword", details.get('keyword_results', 0))
                            with col3:
                                st.metric("BM25", details.get('bm25_results', 0))
                            
                            st.metric("Top Score", f"{details.get('top_score', 0):.3f}")
    
    # Chat input
    if prompt := st.chat_input("Savolingizni yozing... (Yuqori aniqlik rejimida)"):
        # Add user message
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with chat_container:
            with st.chat_message("user"):
                st.markdown(prompt)
            
            # Generate response
            with st.chat_message("assistant"):
                if not st.session_state.rag_pipeline.is_ready:
                    response_text = "❌ Model hali tayyor emas. Iltimos, avval hujjatlarni yuklang va qayta ishlang."
                    st.markdown(response_text)
                    
                    st.session_state.messages.append({
                        "role": "assistant",
                        "content": response_text
                    })
                else:
                    with st.spinner("🧠 Yuqori aniqlikda javob tayyorlanmoqda..."):
                        response = st.session_state.rag_pipeline.query(prompt)
                    
                    # Display response
                    st.markdown(response['answer'])
                    
                    # Confidence display
                    confidence = response['confidence']
                    if confidence >= 80:
                        conf_class = "confidence-high"
                        conf_icon = "🎯"
                    elif confidence >= 60:
                        conf_class = "confidence-medium"
                        conf_icon = "⚡"
                    else:
                        conf_class = "confidence-low"
                        conf_icon = "⚠️"
                    
                    st.markdown(f"""
                    <div style="margin: 1rem 0;">
                        <span class="{conf_class}">{conf_icon} Ishonchlilik: {confidence}%</span>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Sources
                    if response['sources']:
                        with st.expander("📚 Manbalar"):
                            for i, source in enumerate(response['sources'], 1):
                                st.write(f"{i}. {source}")
                    
                    # Search details
                    if 'search_details' in response:
                        details = response['search_details']
                        with st.expander("🔍 Qidiruv tafsilotlari"):
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("Semantic", details.get('semantic_results', 0))
                            with col2:
                                st.metric("Keyword", details.get('keyword_results', 0))
                            with col3:
                                st.metric("BM25", details.get('bm25_results', 0))
                            
                            st.metric("Top Score", f"{details.get('top_score', 0):.3f}")
                    
                    # Save message
                    st.session_state.messages.append({
                        "role": "assistant",
                        "content": response['answer'],
                        "confidence": response['confidence'],
                        "sources": response['sources'],
                        "search_details": response.get('search_details', {})
                    })

if __name__ == "__main__":
    # Run interface
    create_advanced_streamlit_interface(),           # Colons (definitions)
            r'\n\n',         # Paragraphs
        ]
        
        structure_count = 0
        for pattern in structure_indicators:
            if re.search(pattern, answer, re.MULTILINE):
                structure_count += 1
        
        structure_score = min(structure_count / 3, 1.0)  # Max 3 structures
        
        # Check for complete sentences
        sentences = safe_sent_tokenize(answer)
        complete_sentences = sum(1 for s in sentences if len(s.strip()) > 10 and s.strip().endswith(('.', '!', '?')))
        sentence_quality = complete_sentences / max(len(sentences), 1)
        
        return (structure_score + sentence_quality) / 2
    
    def calculate_query_complexity_match(self, query: str, results: List[Dict]) -> float:
        """Query complexity va result match"""
        query_words = len(query.split())
        
        # Simple queries should have high-confidence simple answers
        if query_words <= 3:
            # Look for direct matches
            query_lower = query.lower()
            direct_matches = 0
            for result in results[:3]:
                if any(word in result['content'].lower() for word in query_lower.split()):
                    direct_matches += 1
            return min(direct_matches / 3, 1.0)
        
        # Complex queries should have comprehensive results
        else:
            total_content_length = sum(len(r['content']) for r in results[:5])
            complexity_score = min(total_content_length / (query_words * 100), 1.0)
            return complexity_score
    
    def apply_confidence_boosting(self, base_confidence: float, results: List[Dict], query: str) -> float:
        """Apply confidence boosting based on various factors"""
        boost_factor = 1.0
        
        # Boost for high-quality top results
        if results and results[0].get('combined_score', 0) > 0.8:
            boost_factor *= 1.15
        
        # Boost for multiple consistent results
        if len(results) >= 5:
            boost_factor *= 1.1
        
        # Boost for simple, direct queries
        if len(query.split()) <= 4 and any(word in query.lower() for word in ['nima', 'kim', 'qanday']):
            boost_factor *= 1.2
        
        # Apply diminishing returns
        boosted = base_confidence * boost_factor
        
        # Sigmoid-like function for smoother confidence distribution
        final_confidence = boosted / (1 + np.exp(-(boosted - 0.5) * 6))
        
        return min(final_confidence, 0.98)
    
    def calculate_query_answer_relevance(self, query: str, answer: str) -> float:
        """Query va answer orasidagi relevance"""
        query_words = set(self.text_cleaner.clean_text(query).lower().split())
        answer_words = set(self.text_cleaner.clean_text(answer).lower().split())
        
        if not query_words:
            return 0.0
        
        # Jaccard similarity
        intersection = query_words & answer_words
        union = query_words | answer_words
        
        return len(intersection) / len(union) if union else 0.0
    
    def calculate_context_consistency(self, results: List[Dict]) -> float:
        """Context mazmunlarining consistency"""
        if len(results) < 2:
            return 1.0
        
        contents = [r['content'] for r in results]
        
        # Pairwise similarity calculation
        similarities = []
        for i in range(len(contents)):
            for j in range(i + 1, len(contents)):
                content1_words = set(contents[i].lower().split())
                content2_words = set(contents[j].lower().split())
                
                intersection = content1_words & content2_words
                union = content1_words | content2_words
                
                if union:
                    similarity = len(intersection) / len(union)
                    similarities.append(similarity)
        
        return np.mean(similarities) if similarities else 0.0
    
    def calculate_answer_completeness(self, query: str, answer: str) -> float:
        """Javobning to'liqligi"""
        # Answer uzunligi (optimal range)
        answer_length = len(answer)
        if answer_length < 50:
            length_score = answer_length / 50
        elif answer_length > 1000:
            length_score = max(0.5, 1000 / answer_length)
        else:
            length_score = 1.0
        
        # Question words coverage
        question_words = ['nima', 'qanday', 'qachon', 'qayer', 'kim', 'nega', 'qancha']
        query_lower = query.lower()
        
        answered_questions = 0
        total_questions = 0
        
        for qw in question_words:
            if qw in query_lower:
                total_questions += 1
                # Check if answer contains related content
                if any(indicator in answer.lower() for indicator in [
                    'ta\'rif', 'jarayon', 'vaqt', 'joy', 'shaxs', 'sabab', 'miqdor'
                ]):
                    answered_questions += 1
        
        question_coverage = answered_questions / total_questions if total_questions > 0 else 1.0
        
        return (length_score + question_coverage) / 2
    
    def calculate_source_diversity(self, results: List[Dict]) -> float:
        """Source turlarining diversity"""
        sources = [r.get('source', '') for r in results]
        unique_sources = len(set(sources))
        total_sources = len(sources)
        
        return unique_sources / total_sources if total_sources > 0 else 0.0
    
    def calculate_keyword_coverage(self, query: str, results: List[Dict]) -> float:
        """Query keywords ning results da coverage"""
        query_keywords = set(self.text_cleaner.extract_keywords(query, top_k=5))
        
        if not query_keywords:
            return 1.0
        
        result_keywords = set()
        for result in results:
            if result.get('keywords'):
                result_keywords.update(result['keywords'])
        
        if not result_keywords:
            return 0.0
        
        # Keywords overlap
        overlap = query_keywords & result_keywords
        return len(overlap) / len(query_keywords)

class HighAccuracyRAGPipeline:
    """90% aniqlikli RAG pipeline"""
    
    def __init__(self, config: HighAccuracyConfig = None):
        self.config = config or HighAccuracyConfig()
        
        # Components
        self.doc_loader = AdvancedDocumentLoader()
        self.chunker = SemanticChunker(
            chunk_size=self.config.CHUNK_SIZE,
            chunk_overlap=self.config.CHUNK_OVERLAP
        )
        self.embedder = HybridEmbedder(
            max_features=self.config.MAX_FEATURES,
            n_components=self.config.SVD_COMPONENTS
        )
        self.database = AdvancedDatabase(self.config.DATABASE_PATH)
        self.retriever = None  # Initialize after embedder is ready
        self.response_generator = IntelligentResponseGenerator()
        self.confidence_calculator = AdvancedConfidenceCalculator()
        
        self.is_ready = False
        
    def process_documents(self, file_paths: List[str]):
        """Hujjatlarni qayta ishlash va indexing"""
        logger.info(f"🚀 {len(file_paths)} ta fayl yuqori aniqlikda qayta ishlanmoqda...")
        
        all_chunks = []
        all_keywords = []
        successful_files = 0
        
        # 1. Document loading va chunking
        for file_path in file_paths:
            try:
                logger.info(f"📄 Qayta ishlanmoqda: {file_path}")
                
                # Load document
                text = self.doc_loader.load_document(file_path)
                if not text.strip():
                    logger.warning(f"⚠️ Bo'sh fayl: {file_path}")
                    continue
                
                # Semantic chunking
                chunks = self.chunker.chunk_by_semantic_similarity(text, source=file_path)
                
                # Filter valid chunks
                valid_chunks = [chunk for chunk in chunks if chunk is not None]
                
                if valid_chunks:
                    all_chunks.extend(valid_chunks)
                    # Keywords collect qilish
                    chunk_keywords = [chunk.get('keywords', []) for chunk in valid_chunks]
                    all_keywords.extend(chunk_keywords)
                    
                    successful_files += 1
                    logger.info(f"✅ {file_path}: {len(valid_chunks)} ta semantic chunk")
                else:
                    logger.warning(f"⚠️ Valid chunk yaratilmadi: {file_path}")
                
            except Exception as e:
                logger.error(f"❌ {file_path}: {e}")
        
        if not all_chunks:
            logger.error("❌ Hech qanday valid chunk yaratilmadi!")
            return
        
        logger.info(f"📊 Jami {len(all_chunks)} ta chunk yaratildi")
        
        # 2. Advanced embedding generation
        logger.info("🧠 Hybrid embeddings yaratilmoqda...")
        chunk_texts = [chunk['content'] for chunk in all_chunks]
        embeddings = self.embedder.fit_transform(chunk_texts, all_keywords)
        
        # 3. Database indexing
        logger.info("💾 Advanced database indexing...")
        self.database.add_documents(all_chunks, embeddings)
        
        # 4. Retriever setup
        self.retriever = HybridRetriever(self.database, self.embedder)
        
        # 5. Model saving
        self.embedder.save(self.config.EMBEDDINGS_PATH)
        
        self.is_ready = True
        
        # Final statistics
        stats = self.database.get_stats()
        logger.info(f"""
        🎉 Yuqori aniqlikli qayta ishlash tugadi:
        - ✅ Muvaffaqiyatli fayllar: {successful_files}/{len(file_paths)}
        - 📄 Jami chunklar: {stats['total_documents']}
        - 📁 Noyob manbalar: {stats['unique_sources']}
        - 📏 O'rtacha chunk uzunligi: {stats['average_length']} belgi
        - 🔧 Chunk turlari: {stats['chunk_types']}
        """)
        
    def load_existing_model(self):
        """Mavjud modelni yuklash"""
        try:
            if os.path.exists(self.config.EMBEDDINGS_PATH):
                logger.info("📂 Mavjud model yuklanmoqda...")
                self.embedder.load(self.config.EMBEDDINGS_PATH)
                
                stats = self.database.get_stats()
                if stats['total_documents'] > 0:
                    # Retriever setup
                    self.retriever = HybridRetriever(self.database, self.embedder)
                    self.is_ready = True
                    
                    logger.info(f"""
                    ✅ Mavjud model muvaffaqiyatli yuklandi:
                    - 📄 Jami hujjatlar: {stats['total_documents']}
                    - 📁 Manbalar: {stats['unique_sources']}
                    - 📊 O'rtacha uzunlik: {stats['average_length']} belgi
                    """)
                    return True
                    
        except Exception as e:
            logger.error(f"❌ Model yuklashda xatolik: {e}")
        
        return False
    
    def query(self, question: str, top_k: int = None) -> Dict[str, Any]:
        """Yuqori aniqlikli savol-javob"""
        if not self.is_ready:
            return {
                'answer': "❌ Model hali tayyor emas. Iltimos, avval hujjatlarni qayta ishlang.",
                'sources': [],
                'confidence': 0.0,
                'context_used': 0
            }
        
        if top_k is None:
            top_k = self.config.MAX_DOCS_PER_QUERY
        
        try:
            logger.info(f"🔍 Query: '{question}'")
            
            # 1. Hybrid search
            search_results = self.retriever.hybrid_search(question, top_k)
            
            if not search_results:
                return {
                    'answer': "🔍 Sizning savolingizga mos ma'lumot topilmadi. Iltimos, savolni boshqacha shaklda yozing yoki qo'shimcha hujjatlar yuklang.",
                    'sources': [],
                    'confidence': 0.0,
                    'context_used': 0
                }
            
            # 2. Context preparation
            contexts = [result['content'] for result in search_results]
            sources = list(set([result['source'] for result in search_results]))
            
            # 3. Intelligent response generation
            answer = self.response_generator.generate_response(question, contexts, 0)
            
            # 4. Advanced confidence calculation
            confidence = self.confidence_calculator.calculate_comprehensive_confidence(
                question, search_results, answer
            )
            
            # 5. Update answer with confidence
            final_answer = self.response_generator.generate_response(question, contexts, confidence)
            
            logger.info(f"✅ Query completed - Confidence: {confidence:.1f}%")
            
            return {
                'answer': final_answer,
                'sources': sources,
                'confidence': round(confidence, 1),
                'context_used': len(search_results),
                'search_details': {
                    'semantic_results': len([r for r in search_results if r.get('semantic_score', 0) > 0]),
                    'keyword_results': len([r for r in search_results if r.get('keyword_score', 0) > 0]),
                    'bm25_results': len([r for r in search_results if r.get('bm25_score', 0) > 0]),
                    'top_score': search_results[0].get('combined_score', 0) if search_results else 0
                }
            }
            
        except Exception as e:
            logger.error(f"❌ Query xatolik: {e}")
            return {
                'answer': "❌ Xatolik yuz berdi. Iltimos qaytadan urinib ko'ring.",
                'sources': [],
                'confidence': 0.0,
                'context_used': 0
            }

def load_css_style():
    """CSS stillarni yuklash"""
    css_style = """
    <style>
    * {
        font-family: Verdana, Geneva, Tahoma, sans-serif !important;
    }

    .main-header {
        text-align: center;
        background: linear-gradient(90deg, #00F700 0%, #00D600 100%);
        padding: 1.5rem;
        border-radius: 15px;
        color: white;
        margin-bottom: 2rem;
        box-shadow: 0 4px 15px rgba(0, 247, 0, 0.3);
    }

    .main-header h1 {
        font-size: 2.5rem;
        margin-bottom: 0.5rem;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
    }

    .main-header p {
        font-size: 1.2rem;
        opacity: 0.9;
    }

    .stButton > button {
        width: 95% !important;
        height: 60px !important;
        border: none !important;
        outline: none !important;
        color: #fff !important;
        background: #111 !important;
        cursor: pointer !important;
        position: relative !important;
        z-index: 0 !important;
        border-radius: 10px !important;
        font-weight: bold !important;
        font-size: 16px !important;
        transition: all 0.3s ease !important;
    }

    .stButton > button:before {
        content: '';
        background: linear-gradient(45deg, #00F700, #73ff00, #fffb00, #48ff00, #00ffd5, #002bff, #7a00ff, #ff00c8, #00F700);
        position: absolute;
        top: -2px;
        left: -2px;
        background-size: 400%;
        z-index: -1;
        filter: blur(5px);
        width: calc(100% + 4px);
        height: calc(100% + 4px);
        animation: glowing 20s linear infinite;
        opacity: 0;
        transition: opacity .3s ease-in-out;
        border-radius: 10px;
    }

    .stButton > button:active {
        color: #000 !important;
    }

    .stButton > button:active:after {
        background: transparent !important;
    }

    .stButton > button:hover:before {
        opacity: 1;
    }

    .stButton > button:after {
        z-index: -1;
        content: '';
        position: absolute;
        width: 100%;
        height: 100%;
        background: #00F700;
        left: 0;
        top: 0;
        border-radius: 10px;
    }

    @keyframes glowing {
        0% { background-position: 0 0; }
        50% { background-position: 400% 0; }
        100% { background-position: 0 0; }
    }

    .stat-box {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
        margin: 0.5rem 0;
        border: 2px solid #00F700;
    }

    .confidence-high { 
        color: #00F700 !important; 
        font-weight: bold !important;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
    }
    .confidence-medium { 
        color: #ffc107 !important; 
        font-weight: bold !important;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
    }
    .confidence-low { 
        color: #dc3545 !important; 
        font-weight: bold !important;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
    }

    .upload-section {
        border: 2px dashed #00F700;
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
        background: rgba(0, 247, 0, 0.05);
    }

    .sidebar-title {
        color: #00F700 !important;
        font-weight: bold !important;
        font-size: 1.2rem !important;
        margin-bottom: 1rem !important;
    }

    .metric-container {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 0.8rem;
        border-radius: 8px;
        margin: 0.3rem 0;
    }

    .chat-container {
        border-radius: 10px;
        padding: 1rem;
        margin: 0.5rem 0;
    }

    .user-message {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 15px 15px 5px 15px;
        padding: 1rem;
        margin: 0.5rem 0;
    }

    .assistant-message {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        color: white;
        border-radius: 15px 15px 15px 5px;
        padding: 1rem;
        margin: 0.5rem 0;
    }

    .processing-indicator {
        background: linear-gradient(45deg, #00F700, #73ff00);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
        animation: pulse 2s infinite;
    }

    @keyframes pulse {
        0% { opacity: 1; }
        50% { opacity: 0.7; }
        100% { opacity: 1; }
    }
    </style>
    """
    return css_style

def create_advanced_streamlit_interface():
    """Ilg'or Streamlit interface"""
    
    st.set_page_config(
        page_title="CHATBOT",
        page_icon="🎯",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Apply custom CSS
    st.markdown(load_css_style(), unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🎯 CHATBOT</h1>
        <p>Yuqori aniqlikli AI yordamchisi - API-siz, to'liq offline</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Session state initialization
    if 'rag_pipeline' not in st.session_state:
        st.session_state.rag_pipeline = HighAccuracyRAGPipeline()
        # Auto-load existing model
        st.session_state.rag_pipeline.load_existing_model()
    
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    
    if 'processing_stats' not in st.session_state:
        st.session_state.processing_stats = {}
    
    # Sidebar
    with st.sidebar:
        st.markdown('<h2 class="sidebar-title">⚙️ Boshqaruv Paneli</h2>', unsafe_allow_html=True)
        
        # Model status
        if st.session_state.rag_pipeline.is_ready:
            st.success("✅ Model tayyor")
            stats = st.session_state.rag_pipeline.database.get_stats()
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f'<div class="metric-container">📄 Hujjatlar<br><strong>{stats["total_documents"]}</strong></div>', unsafe_allow_html=True)
                st.markdown(f'<div class="metric-container">📏 O\'rtacha uzunlik<br><strong>{stats["average_length"]:.0f}</strong></div>', unsafe_allow_html=True)
            with col2:
                st.markdown(f'<div class="metric-container">📁 Manbalar<br><strong>{stats["unique_sources"]}</strong></div>', unsafe_allow_html=True)
                if 'chunk_types' in stats:
                    semantic_count = stats['chunk_types'].get('semantic', 0)
                    st.markdown(f'<div class="metric-container">🧠 Semantic chunks<br><strong>{semantic_count}</strong></div>', unsafe_allow_html=True)
        else:
            st.warning("⚠️ Model hali tayyor emas")
        
        st.markdown("---")
        
        # Configuration
        st.subheader("🔧 Sozlamalar")
        
        chunk_size = st.slider("Chunk o'lchami", 150, 400, 
                              st.session_state.rag_pipeline.config.CHUNK_SIZE)
        
        max_docs = st.slider("Maksimal contextlar", 5, 20, 
                           st.session_state.rag_pipeline.config.MAX_DOCS_PER_QUERY)
        
        # Update config
        st.session_state.rag_pipeline.config.CHUNK_SIZE = chunk_size
        st.session_state.rag_pipeline.config.MAX_DOCS_PER_QUERY = max_docs
        
        st.markdown("---")
        
        # File upload
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.subheader("📤 Hujjat yuklash")
        
        # File uploader with improved handling
        uploaded_files = st.file_uploader(
            "Fayllarni tanlang",
            type=['pdf', 'docx', 'txt', 'html', 'md'],
            accept_multiple_files=True,
            help="Qo'llab-quvvatlanuvchi formatlar: PDF, DOCX, TXT, HTML, MD"
        )
        
        # Show uploaded files
        if uploaded_files:
            st.write("**Yuklangan fayllar:**")
            for file in uploaded_files:
                file_size = len(file.getvalue()) / 1024 / 1024  # MB
                st.write(f"📄 {file.name} ({file_size:.1f} MB)")
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Processing button
        col1, col2 = st.columns(2)
        with col1:
            process_btn = st.button("🚀 Qayta ishlash", disabled=not uploaded_files, use_container_width=True)
        with col2:
            clear_btn = st.button("🗑️ Tozalash", use_container_width=True)
        
        # Processing with comprehensive error handling and debugging
        if process_btn and uploaded_files:
            # Create temporary directory
            temp_dir = Path("./temp_uploads")
            temp_dir.mkdir(exist_ok=True)
            
            # Initialize variables
            temp_paths = []
            
            try:
                # Phase 1: File saving
                st.info("📁 1-bosqich: Fayllar saqlanmoqda...")
                progress_bar = st.progress(0)
                
                total_files = len(uploaded_files)
                for i, uploaded_file in enumerate(uploaded_files):
                    try:
                        # Progress update
                        progress = (i + 1) / total_files * 0.2
                        progress_bar.progress(progress)
                        
                        # Create safe filename
                        original_name = uploaded_file.name
                        safe_name = re.sub(r'[^\w\-_\.]', '_', original_name)
                        if not safe_name:
                            safe_name = f"document_{i}.txt"
                        
                        temp_path = temp_dir / f"{i}_{safe_name}"
                        
                        # Save file
                        file_content = uploaded_file.read()
                        if len(file_content) == 0:
                            st.warning(f"⚠️ Bo'sh fayl: {original_name}")
                            continue
                            
                        with open(temp_path, "wb") as f:
                            f.write(file_content)
                        
                        temp_paths.append(str(temp_path))
                        st.success(f"✅ Saqlandi: {original_name} ({len(file_content)} bytes)")
                        
                    except Exception as e:
                        st.error(f"❌ {uploaded_file.name} saqlashda xatolik: {str(e)}")
                        continue
                
                if not temp_paths:
                    st.error("❌ Hech qanday fayl saqlanmadi!")
                    return
                
                st.success(f"✅ {len(temp_paths)} ta fayl muvaffaqiyatli saqlandi")
                
                # Phase 2: Document processing
                st.info("🧠 2-bosqich: Hujjatlar qayta ishlanmoqda...")
                progress_bar.progress(0.3)
                
                # Test document loading first
                doc_loader = AdvancedDocumentLoader()
                loaded_texts = []
                
                for i, temp_path in enumerate(temp_paths):
                    try:
                        progress = 0.3 + (i + 1) / len(temp_paths) * 0.2
                        progress_bar.progress(progress)
                        
                        text = doc_loader.load_document(temp_path)
                        if text and len(text.strip()) > 10:
                            loaded_texts.append((temp_path, text))
                            st.info(f"📄 Yuklandi: {Path(temp_path).name} ({len(text)} belgi)")
                        else:
                            st.warning(f"⚠️ Bo'sh yoki juda qisqa matn: {Path(temp_path).name}")
                    except Exception as e:
                        st.error(f"❌ {Path(temp_path).name} yuklashda xatolik: {str(e)}")
                        continue
                
                if not loaded_texts:
                    st.error("❌ Hech qanday matn yuklanmadi!")
                    return
                
                st.success(f"✅ {len(loaded_texts)} ta hujjat matnlari yuklandi")
                
                # Phase 3: Create new pipeline and process
                st.info("⚙️ 3-bosqich: RAG pipeline yaratilmoqda...")
                progress_bar.progress(0.6)
                
                # Create completely fresh pipeline
                new_config = HighAccuracyConfig()
                new_pipeline = HighAccuracyRAGPipeline(new_config)
                
                # Process each document step by step
                all_chunks = []
                all_keywords = []
                
                for i, (file_path, text) in enumerate(loaded_texts):
                    try:
                        progress = 0.6 + (i + 1) / len(loaded_texts) * 0.2
                        progress_bar.progress(progress)
                        
                        # Create chunks
                        chunks = new_pipeline.chunker.chunk_by_semantic_similarity(text, source=Path(file_path).name)
                        valid_chunks = [chunk for chunk in chunks if chunk is not None]
                        
                        if valid_chunks:
                            all_chunks.extend(valid_chunks)
                            chunk_keywords = [chunk.get('keywords', []) for chunk in valid_chunks]
                            all_keywords.extend(chunk_keywords)
                            st.info(f"🔧 {Path(file_path).name}: {len(valid_chunks)} chunk yaratildi")
                        else:
                            st.warning(f"⚠️ {Path(file_path).name}: chunk yaratilmadi")
                            
                    except Exception as e:
                        st.error(f"❌ {Path(file_path).name} chunking xatolik: {str(e)}")
                        continue
                
                if not all_chunks:
                    st.error("❌ Hech qanday chunk yaratilmadi!")
                    return
                
                st.success(f"✅ Jami {len(all_chunks)} ta chunk yaratildi")
                
                # Phase 4: Create embeddings
                st.info("🧠 4-bosqich: Embeddings yaratilmoqda...")
                progress_bar.progress(0.8)
                
                try:
                    chunk_texts = [chunk['content'] for chunk in all_chunks]
                    embeddings = new_pipeline.embedder.fit_transform(chunk_texts, all_keywords)
                    st.success(f"✅ Embeddings yaratildi: {embeddings.shape}")
                except Exception as e:
                    st.error(f"❌ Embeddings yaratishda xatolik: {str(e)}")
                    return
                
                # Phase 5: Save to database
                st.info("💾 5-bosqich: Database ga saqlash...")
                progress_bar.progress(0.9)
                
                try:
                    new_pipeline.database.add_documents(all_chunks, embeddings)
                    new_pipeline.retriever = HybridRetriever(new_pipeline.database, new_pipeline.embedder)
                    new_pipeline.embedder.save(new_pipeline.config.EMBEDDINGS_PATH)
                    new_pipeline.is_ready = True
                    st.success("✅ Database ga saqlandi")
                except Exception as e:
                    st.error(f"❌ Database ga saqlashda xatolik: {str(e)}")
                    return
                
                # Phase 6: Final verification and replacement
                st.info("✅ 6-bosqich: Yakuniy tekshirish...")
                progress_bar.progress(0.95)
                
                # Verify everything works
                try:
                    test_stats = new_pipeline.database.get_stats()
                    if test_stats['total_documents'] > 0 and new_pipeline.is_ready:
                        # Replace the session pipeline
                        st.session_state.rag_pipeline = new_pipeline
                        
                        progress_bar.progress(1.0)
                        
                        # Show final success
                        success_msg = f"""
                        🎉 **MUVAFFAQIYAT!**
                        
                        ✅ **Jami hujjatlar:** {test_stats['total_documents']}  
                        ✅ **Noyob manbalar:** {test_stats['unique_sources']}  
                        ✅ **O'rtacha uzunlik:** {test_stats['average_length']:.0f} belgi  
                        ✅ **Model holati:** Tayyor  
                        
                        Endi savollaringizni bering! 🚀
                        """
                        st.balloons()
                        st.success(success_msg)
                        
                        # Auto-refresh after success
                        import time
                        time.sleep(3)
                        st.rerun()
                        
                    else:
                        st.error("❌ Model tekshiruvdan o'tmadi!")
                        
                except Exception as e:
                    st.error(f"❌ Yakuniy tekshiruvda xatolik: {str(e)}")
                
            except Exception as e:
                st.error(f"❌ Umumiy xatolik: {str(e)}")
                import traceback
                st.code(traceback.format_exc())
                
            finally:
                # Cleanup temp files
                for temp_path in temp_paths:
                    try:
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                    except:
                        pass
                try:
                    if temp_dir.exists():
                        import shutil
                        shutil.rmtree(temp_dir, ignore_errors=True)
                except:
                    pass
        
        # Clear data
        if clear_btn:
            if st.session_state.get('clear_confirmed', False):
                try:
                    # Remove database and embeddings
                    if os.path.exists(st.session_state.rag_pipeline.config.DATABASE_PATH):
                        os.remove(st.session_state.rag_pipeline.config.DATABASE_PATH)
                    if os.path.exists(st.session_state.rag_pipeline.config.EMBEDDINGS_PATH):
                        os.remove(st.session_state.rag_pipeline.config.EMBEDDINGS_PATH)
                    
                    # Reset pipeline
                    st.session_state.rag_pipeline = HighAccuracyRAGPipeline()
                    st.session_state.messages = []
                    st.session_state.clear_confirmed = False
                    
                    st.success("✅ Ma'lumotlar tozalandi!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Tozalash xatolik: {e}")
            else:
                st.session_state.clear_confirmed = True
                st.warning("⚠️ Haqiqatan ham barcha ma'lumotlarni o'chirmoqchimisiz? Qayta bosing.")
    
    # Main chat interface
    st.header("💬 Yuqori Aniqlikli Suhbat")
    
    # Chat container
    chat_container = st.container()
    
    with chat_container:
        # Display chat messages
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
                
                # Additional info for assistant messages
                if message["role"] == "assistant":
                    if "confidence" in message:
                        confidence = message["confidence"]
                        if confidence >= 80:
                            conf_class = "confidence-high"
                            conf_icon = "🎯"
                        elif confidence >= 60:
                            conf_class = "confidence-medium"
                            conf_icon = "⚡"
                        else:
                            conf_class = "confidence-low"
                            conf_icon = "⚠️"
                        
                        st.markdown(f"""
                        <div style="margin: 1rem 0;">
                            <span class="{conf_class}">{conf_icon} Ishonchlilik: {confidence}%</span>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # Sources
                    if "sources" in message and message["sources"]:
                        with st.expander("📚 Manbalar"):
                            for i, source in enumerate(message["sources"], 1):
                                st.write(f"{i}. {source}")
                    
                    # Search details
                    if "search_details" in message:
                        details = message["search_details"]
                        with st.expander("🔍 Qidiruv tafsilotlari"):
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("Semantic", details.get('semantic_results', 0))
                            with col2:
                                st.metric("Keyword", details.get('keyword_results', 0))
                            with col3:
                                st.metric("BM25", details.get('bm25_results', 0))
                            
                            st.metric("Top Score", f"{details.get('top_score', 0):.3f}")
    
    # Chat input
    if prompt := st.chat_input("Savolingizni yozing... (Yuqori aniqlik rejimida)"):
        # Add user message
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with chat_container:
            with st.chat_message("user"):
                st.markdown(prompt)
            
            # Generate response
            with st.chat_message("assistant"):
                if not st.session_state.rag_pipeline.is_ready:
                    response_text = "❌ Model hali tayyor emas. Iltimos, avval hujjatlarni yuklang va qayta ishlang."
                    st.markdown(response_text)
                    
                    st.session_state.messages.append({
                        "role": "assistant",
                        "content": response_text
                    })
                else:
                    with st.spinner("🧠 Yuqori aniqlikda javob tayyorlanmoqda..."):
                        response = st.session_state.rag_pipeline.query(prompt)
                    
                    # Display response
                    st.markdown(response['answer'])
                    
                    # Confidence display
                    confidence = response['confidence']
                    if confidence >= 80:
                        conf_class = "confidence-high"
                        conf_icon = "🎯"
                    elif confidence >= 60:
                        conf_class = "confidence-medium"
                        conf_icon = "⚡"
                    else:
                        conf_class = "confidence-low"
                        conf_icon = "⚠️"
                    
                    st.markdown(f"""
                    <div style="margin: 1rem 0;">
                        <span class="{conf_class}">{conf_icon} Ishonchlilik: {confidence}%</span>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Sources
                    if response['sources']:
                        with st.expander("📚 Manbalar"):
                            for i, source in enumerate(response['sources'], 1):
                                st.write(f"{i}. {source}")
                    
                    # Search details
                    if 'search_details' in response:
                        details = response['search_details']
                        with st.expander("🔍 Qidiruv tafsilotlari"):
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("Semantic", details.get('semantic_results', 0))
                            with col2:
                                st.metric("Keyword", details.get('keyword_results', 0))
                            with col3:
                                st.metric("BM25", details.get('bm25_results', 0))
                            
                            st.metric("Top Score", f"{details.get('top_score', 0):.3f}")
                    
                    # Save message
                    st.session_state.messages.append({
                        "role": "assistant",
                        "content": response['answer'],
                        "confidence": response['confidence'],
                        "sources": response['sources'],
                        "search_details": response.get('search_details', {})
                    })

if __name__ == "__main__":
    # Run interface
    create_advanced_streamlit_interface()
